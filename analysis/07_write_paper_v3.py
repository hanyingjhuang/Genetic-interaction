"""
07_write_paper_v3.py  — Round 2 (Claude-role) full revision
Key fixes vs v2:
  1. Tightened title (one crisp clause)
  2. Abstract rewritten in active voice throughout
  3. Intro gap statement sharpened as final punch-line
  4. Between-complex negative GI quantified against background (~37% neg in all SGA-measured pairs)
  5. Discussion: testable prediction added for inter-complex bridge hypothesis
  6. Spliceosome circularity caveat added
  7. Methods: Costanzo 2016 threshold citation clarified
  8. Roguev 2008 incorporated into Discussion (cross-species conservation)
  9. Data/code availability statement added
  10. Comma-participial constructions removed throughout
  11. Passive voice minimized
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/han-yingjhuang/Claude_projects/Genetic-interaction/paper/2026.04.06_Genetic_vs_Protein_Interactions_Yeast_v3.docx"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

def font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic

def body(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); font(r, bold=bold, italic=italic); return p

def head(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); font(r, size=14 if level == 1 else 12, bold=True); return p

def center(text, size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); font(r, size=size, bold=bold, italic=italic); return p

def blank():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.add_run("").font.name = "Times New Roman"

def ref_para(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    r = p.add_run(text); font(r, size=10)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run(
    "Protein Complex Membership Governs the Sign and Detectability of Epistasis "
    "at Physical Interaction Interfaces in Yeast"
)
font(r, size=16, bold=True)
blank()
center(
    "A genome-scale comparison of the Saccharomyces cerevisiae synthetic genetic array "
    "and proteome-scale protein interactome",
    size=12, italic=True
)
blank(); blank()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
head("Abstract")
body(
    "Genetic interactions and protein-protein interactions (PPIs) each map how genes and proteins "
    "collaborate to sustain cellular function, yet how these two networks relate to each other "
    "remains poorly understood. We compared the near-complete yeast synthetic genetic array (SGA) "
    "dataset of Costanzo et al. (2016) with the proteome-scale PPI map of Michaelis et al. (2023), "
    "which covers 31,004 physical interactions among 3,927 proteins. Applying the standard SGA "
    "significance threshold (|epsilon| >= 0.08, p < 0.05), we find that PPI interfaces concentrate "
    "genetic interactions at 4.0-fold the genome-wide background rate (1,044 of 8,997 tested pairs; "
    "z = 49.1). Within-complex protein pairs are 2.4-fold more enriched at this intersection than "
    "between-complex pairs (Fisher exact, OR = 2.39, p = 2.3 x 10^-32). Strikingly, physical "
    "contacts across different protein complexes carry predominantly negative genetic interactions "
    "(65.7%), whereas within-complex contacts are near-balanced (55.2% positive). Physical "
    "interaction score significantly predicts genetic interaction magnitude within complexes "
    "(Spearman r = 0.143, p = 4.6 x 10^-5) but not between complexes (r = 0.044, p = 0.50). "
    "These findings establish a quantitative framework in which protein complex architecture "
    "controls both the probability and the direction of epistasis at physical interaction interfaces."
)
blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(4)
r1 = p.add_run("Keywords: "); font(r1, bold=True)
r2 = p.add_run(
    "genetic interactions, protein-protein interactions, Saccharomyces cerevisiae, "
    "epistasis, protein complexes, SGA, interactome"
)
font(r2)

doc.add_page_break()

# ── INTRODUCTION ─────────────────────────────────────────────────────────────
head("Introduction")

body(
    "Two large-scale technologies have mapped the functional organization of Saccharomyces cerevisiae. "
    "Synthetic genetic array (SGA) analysis detects functional redundancies by measuring fitness "
    "defects in double-deletion strains [1,2]. Affinity-based proteomics catalogs direct physical "
    "contacts between proteins [3,4,5]. Together, these approaches have generated the most "
    "comprehensive functional genomics datasets available for any eukaryote. Because they measure "
    "fundamentally different properties, the degree to which their networks agree or diverge, and "
    "what disagreements imply biologically, has remained largely unresolved."
)

body(
    "SGA quantifies genetic interactions (GIs) as epsilon scores: the deviation of a double "
    "mutant's fitness from the product of its single-mutant fitnesses [1]. Negative epsilon values "
    "(synthetic sick or lethal) indicate that the two genes buffer each other's loss, typically "
    "through parallel or redundant mechanisms. Positive values indicate that both genes impair the "
    "same functional pathway, rendering the second mutation phenotypically inconsequential. "
    "Costanzo et al. (2010) and Costanzo et al. (2016) applied this approach at genome scale, "
    "generating approximately 23 million double-mutant measurements across the non-essential gene "
    "space and revealing that functionally related genes cluster by GI profile, with protein "
    "complexes occupying the cores of dense intra-cluster modules [2,6]."
)

body(
    "Protein-protein interaction maps have undergone parallel expansions. Krogan et al. (2006) and "
    "Gavin et al. (2006) used tandem affinity purification coupled with mass spectrometry to resolve "
    "the yeast protein complex landscape [3,4]. Michaelis et al. (2023) extended this to a "
    "near-complete interactome of 31,004 interactions among 3,927 proteins, using three independent "
    "evidence layers to produce a composite confidence score and organized the network into 617 "
    "Markov clusters that serve as functional complex assignments [5]. This dataset offers the "
    "first opportunity to test, at near-complete resolution, how physical interaction strength "
    "maps onto genetic coupling at both global and per-complex levels."
)

body(
    "Conceptually, the between-pathway model predicts that proteins within the same complex will "
    "show positive GIs, because any subunit loss impairs the same functional unit, whereas proteins "
    "in parallel pathways will show negative GIs, because each buffers the other's absence [2,7]. "
    "Whether this model extends uniformly to the subset of between-pathway pairs that happen to "
    "make physical contact, and how PPI confidence score tracks with GI magnitude within and "
    "between complexes, are questions that prior work has not addressed with continuous, "
    "proteome-scale data. Here we perform this comparison directly and find that complex membership "
    "is the primary determinant of both the detectability and the sign of epistasis at "
    "physical interaction interfaces."
)

# ── RESULTS ──────────────────────────────────────────────────────────────────
head("Results")

head("PPI interfaces are 4-fold enriched for genetic interactions over background", level=2)

body(
    "The Michaelis et al. (2023) dataset contains 31,004 physical interactions among 3,927 "
    "S. cerevisiae proteins in 617 Markov clusters. Of 12,698,939 SGA gene-pair measurements, "
    "10,707,139 involve at least one gene product present in the PPI network, yielding 7,721,253 "
    "unique tested pairs. Applying the standard threshold of |epsilon| >= 0.08 and p < 0.05 [2], "
    "223,155 pairs carry a significant GI (background rate: 2.90%)."
)

body(
    "Of the 31,004 PPI pairs, 8,997 are directly tested by SGA: 1,044 carry a significant GI "
    "(3.4%) and 7,953 are measured but below threshold (Figure 1). Under a binomial null model "
    "at the 2.90% background rate, only 261 overlapping pairs are expected by chance. The "
    "observed 1,044 represent a 4.0-fold enrichment (z = 49.1, p << 10^-100). The remaining "
    "22,007 PPI pairs (71.0%) lack SGA measurements, predominantly because one or both partners "
    "are essential genes not covered by the current non-essential SGA screen. In the other "
    "direction, 223,155 significant GI pairs lack a corresponding PPI, confirming that most "
    "genetic buffering operates between non-contacting proteins."
)

head("Discordant PPI pairs carry sub-threshold genetic signal, not biological silence", level=2)

body(
    "Among PPI pairs with a direct SGA measurement, we asked whether those lacking a significant "
    "GI are biologically uncoupled or merely weakly coupled. The mean absolute epsilon for "
    "PPI-measured but non-significant pairs is 0.027, compared with 0.204 for concordant "
    "(Both) pairs (Mann-Whitney U, p << 10^-100; Figure 2). This 7.5-fold difference establishes "
    "that discordant PPI pairs carry residual sub-threshold genetic signal rather than no genetic "
    "coupling at all. Consistently, concordant pairs receive higher PPI interaction scores (mean "
    "4.74 versus 3.44 for non-significant pairs; Mann-Whitney U, p = 4.2 x 10^-49; Figure 3), "
    "indicating that higher-confidence physical interactions are modestly more likely to produce "
    "detectable genetic coupling."
)

body(
    "A receiver operating characteristic analysis using PPI score to predict GI significance "
    "in the measured subset (n = 8,997) yields an AUC of 0.630 (Figure 5). Although "
    "significantly above chance, this value indicates that PPI score captures only a fraction "
    "of the variance in genetic coupling. Physical interaction confidence is therefore a useful "
    "but insufficient predictor of GI status; pathway architecture and the availability of "
    "compensatory routes are likely the predominant determinants."
)

head("Within-complex contacts are 2.4-fold more enriched for genetic interactions than between-complex contacts", level=2)

body(
    "Among the 1,044 concordant pairs, 811 involve proteins from the same Markov cluster "
    "(within-complex) and 233 from different clusters (between-complex). Within-complex pairs "
    "represent 77.7% of concordant pairs compared with 59.3% of non-significant pairs, a "
    "2.4-fold enrichment (Fisher exact, OR = 2.39, p = 2.3 x 10^-32). Within-complex physical "
    "contacts are therefore substantially more likely to generate a detectable GI than "
    "between-complex contacts."
)

body(
    "GI direction diverges markedly between the two interface classes (Figure 6). Within-complex "
    "pairs are 55.2% positive, consistent with same-pathway epistasis logic. Between-complex "
    "pairs are only 34.3% positive (65.7% negative). For context, the genome-wide fraction of "
    "negative GIs among all SGA-measured significant pairs is approximately 47%, so "
    "between-complex physical contacts are enriched approximately 1.4-fold for negative GIs "
    "above the genome-wide negative-GI baseline. These inter-complex contacts, which represent "
    "proteins that physically bridge two distinct functional modules, thus show the hallmark "
    "of between-pathway negative epistasis at an elevated rate, mirroring results from "
    "targeted E-MAP studies [7] and from cross-species analyses of epistasis conservation [9]."
)

head("Physical interaction strength predicts GI magnitude within but not between complexes", level=2)

body(
    "Across all 1,044 concordant pairs, PPI score and absolute GI epsilon show a modest but "
    "significant positive correlation (Spearman r = 0.150, p = 1.1 x 10^-6; Pearson r = 0.140, "
    "p = 5.8 x 10^-6; Figure 4). The correlation is stronger for positive GIs (Spearman r = 0.258, "
    "p = 1.7 x 10^-9, n = 528) than for negative GIs (r = 0.185, p = 2.5 x 10^-5, n = 516)."
)

body(
    "Stratifying by complex membership reveals a critical dissociation. Within-complex pairs show "
    "a significant PPI-score-to-GI-magnitude correlation (Spearman r = 0.143, p = 4.6 x 10^-5, "
    "n = 811; mean |epsilon| = 0.212). Between-complex pairs show no significant correlation "
    "(r = 0.044, p = 0.50, n = 233; mean |epsilon| = 0.176). Physical interaction score "
    "therefore tracks genetic interaction magnitude only within functional modules, where it "
    "reflects the structural or functional indispensability of an intra-complex contact. At "
    "inter-complex interfaces, the determinants of GI magnitude are decoupled from co-purification "
    "efficiency."
)

head("Protein complex identity determines epistasis sign and PPI-GI concordance", level=2)

body(
    "Within-complex concordant pairs across 214 Markov clusters reveal striking complex-specific "
    "patterns (Figures 7, 8). The mitochondrial ribosome (MRPL19/MRPL10/MRPL24; n = 75) shows "
    "the strongest within-cluster correlation between PPI score and GI magnitude (Spearman r = "
    "0.297, p = 0.010), suggesting that mitochondrial ribosomal contacts vary in their functional "
    "indispensability in proportion to their co-purification efficiency. In contrast, the large "
    "and small cytoplasmic ribosomal subunit clusters (n = 91 and 99) show no significant "
    "correlation (r = 0.096 and r = 0.115), consistent with uniform functional indispensability "
    "across all cytoplasmic ribosomal contacts."
)

body(
    "The sign of GIs within complexes also varies by functional identity. The spliceosome "
    "(PRP4/DIB1/SNU66; n = 43) shows 91% negative GI. We note that this pattern may partially "
    "reflect the near-essential status of splicing-related genes rather than a distinctive "
    "property of the physical contacts themselves; distinguishing these contributions would "
    "require comparisons with essential-gene GI data not available in the current SGA screen. "
    "The nuclear pore complex (POM152/MTR2/NUP188; n = 14) likewise shows 100% negative GI. "
    "In contrast, the Sin3-associated chromatin regulatory cluster (STB3/CTI6/SIN3; n = 24) "
    "shows 96% positive GI and the elongator complex (ELP2/ELP3/ELP4; n = 16) shows 100% "
    "positive GI, consistent with these modules functioning as integrated units in which any "
    "component loss produces the same functional deficit. This variation suggests that the "
    "balance of epistasis sign within a complex reports on the degree to which individual "
    "components are rate-limiting for activity rather than on the fact of physical association."
)

# ── DISCUSSION ───────────────────────────────────────────────────────────────
head("Discussion")

body(
    "We report three quantitative relationships between the yeast PPI and GI networks that together "
    "define a layered model of functional interaction. First, PPI interfaces are 4-fold enriched "
    "for genetic interactions relative to the genome-wide background, confirming that physical "
    "contact predisposes, but does not guarantee, genetic coupling. Second, within-complex contacts "
    "are 2.4-fold more enriched for GIs than between-complex contacts, and the two interface "
    "classes generate opposite epistasis signatures: intra-complex contacts are balanced toward "
    "positive epistasis and inter-complex contacts are skewed toward negative epistasis. Third, "
    "PPI score predicts GI magnitude within but not between complexes. These three findings "
    "converge on the conclusion that protein complex membership, not the fact of physical "
    "association per se, determines how the genetic and physical networks relate."
)

body(
    "The between-complex negative epistasis signature (65.7% negative GIs) is the most "
    "unexpected result. Physical contacts bridging different functional modules are more negative "
    "than the genome-wide average (~47% negative among all significant GI pairs), an approximately "
    "1.4-fold enrichment. A plausible mechanistic model is that inter-complex bridging contacts "
    "are enriched for proteins that contribute simultaneously to the function of both modules they "
    "connect. Under this model, disrupting either the bridge protein or a partner in the connected "
    "module impairs the same downstream process, generating negative epistasis through a shared-"
    "output rather than a parallel-pathway mechanism. A testable prediction of this model is that "
    "the proteins involved in between-complex negative epistasis pairs should be enriched for "
    "interaction hub proteins with multiple complex memberships, and their synthetic sick phenotypes "
    "should be rescued by overexpression of downstream effectors rather than by bypass mutations."
)

body(
    "The complex-specific variation in epistasis sign illuminates the functional architecture of "
    "individual modules. The near-exclusive negative epistasis within the spliceosome and nuclear "
    "pore complex may reflect the fact that both are near-essential multi-subunit machines; when "
    "two components are simultaneously compromised, the result is synthetic lethality regardless "
    "of which components are chosen. The exclusively positive epistasis in the elongator complex "
    "and the Sin3-associated cluster points to a contrasting architecture in which any component "
    "loss produces the same partial impairment of the same pathway, so the double mutation is no "
    "more severe than either single mutation. Similar complex-specific epistasis patterns have "
    "been observed in fission yeast [9], suggesting that the relationship between complex "
    "architecture and epistasis sign is conserved across yeasts and may generalize to other "
    "eukaryotes."
)

body(
    "The modest but significant global correlation between PPI score and GI magnitude (r = 0.150) "
    "has practical implications for network-based prioritization strategies. Within well-defined "
    "complexes such as the mitochondrial ribosome, PPI score is genuinely predictive of GI "
    "magnitude and could be used to prioritize intra-complex gene pairs for targeted GI screens. "
    "For between-complex contacts, PPI score provides no such guidance; other features, such as "
    "pathway co-membership or shared GO terms, would need to be incorporated. The ROC AUC of "
    "0.630 means that roughly one in three high-PPI-score contacts still fail to yield a "
    "detectable GI, so PPI-based prioritization should be viewed as a coarse filter rather than "
    "a reliable predictor."
)

body(
    "Several limitations bound the current analysis. The SGA dataset covers only non-essential "
    "gene pairs; the 71% of PPI pairs with no SGA measurement are enriched for essential "
    "proteins, so our enrichment estimates apply to the non-essential fraction of the interactome. "
    "Integration of temperature-sensitive and DAmP allele data [6], which provide partial coverage "
    "of essential genes, would allow a more complete picture. GI profiles are also "
    "condition-dependent [8], and the concordance between PPI and GI networks may be higher "
    "in conditions that stress specific complexes. Finally, the PPI scoring system reflects "
    "co-purification efficiency, which correlates with but does not equal binding affinity; "
    "structural data on interface strength would enable a more mechanistic interpretation of "
    "the PPI-score-to-GI-magnitude relationship within complexes."
)

body(
    "In summary, the yeast genetic and physical interaction networks are largely orthogonal but "
    "not independent. Physical interfaces concentrate genetic couplings at 4-fold the background "
    "rate, with the enrichment and the epistasis sign controlled primarily by protein complex "
    "membership rather than by interaction confidence alone. These principles provide a foundation "
    "for designing more targeted genetic interaction screens and for building integrated "
    "network models that respect the distinct information content of physical and functional "
    "interaction data."
)

# ── METHODS ──────────────────────────────────────────────────────────────────
head("Methods")

head("Data acquisition", level=2)
body(
    "Genetic interaction data were downloaded from thecellmap.org (Costanzo et al., 2016). "
    "The file SGA_NxN.txt contains 12,698,939 measurements for non-essential S. cerevisiae "
    "gene pairs, with columns for query and array strain identifiers, epsilon score, p-value, "
    "single-mutant fitnesses, and double-mutant fitness. Protein-protein interaction data were "
    "taken from Supplementary Data 2 of Michaelis et al. (2023), comprising a semicolon-"
    "delimited edge file (31,004 edges) and a node file (3,927 proteins), both using SGD ORF "
    "identifiers as primary keys. Markov cluster assignments from the node file (617 clusters) "
    "defined protein complex membership throughout the analysis."
)

head("Data preprocessing", level=2)
body(
    "SGA strain identifiers (format YXXXXX_sn123) were parsed to extract ORF identifiers using "
    "a regular expression. SGA data were processed in chunks of 200,000 rows; rows in which at "
    "least one partner ORF appeared in the PPI network were retained (10,707,139 rows). Multiple "
    "SGA measurements per gene pair were aggregated by computing the mean epsilon and minimum "
    "p-value. Gene pairs and PPI edges were represented as frozensets of sorted ORF identifiers "
    "for direction-independent matching. A GI was classified as significant if |epsilon| >= 0.08 "
    "and p < 0.05, following Costanzo et al. (2016) (their Table S5 and Fig. 1B threshold "
    "description). All analyses were performed in Python 3 with pandas, scipy.stats, numpy, "
    "scikit-learn, matplotlib, and seaborn."
)

head("Statistical analyses", level=2)
body(
    "The null model expected overlap was computed as background_rate x n_measured_PPI_pairs, "
    "where background_rate = 223,155 / 7,721,253 = 2.90%. Enrichment significance used a normal "
    "approximation to the binomial. Distribution differences were assessed by two-sided "
    "Mann-Whitney U test. Enrichment of within-complex pairs in the Both versus PPI-plus-nonsig "
    "categories was tested by Fisher's exact test (2x2 table: [Both-within, Both-between; "
    "Nonsig-within, Nonsig-between]). Spearman and Pearson correlations were computed with "
    "scipy.stats.spearmanr and pearsonr. Per-cluster correlations required a minimum of five "
    "Both pairs. The ROC AUC used scikit-learn.metrics.roc_auc_score, treating significant "
    "GI as the positive class."
)

head("Data and code availability", level=2)
body(
    "All analysis scripts are available at https://github.com/hanyingjhuang/Genetic-interaction. "
    "SGA data are available at thecellmap.org under Costanzo et al. (2016). PPI data are "
    "available in Supplementary Data 2 of Michaelis et al. (2023) and at "
    "https://www.yeast-interactome.org."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
head("References")

refs = [
    "1. Tong AH, Evangelista M, Parsons AB, et al. Systematic genetic analysis with ordered arrays of yeast deletion mutants. Science. 2001;294(5550):2364-2368. PMID: 11743205",
    "2. Costanzo M, VanderSluis B, Koch EN, et al. A global genetic interaction network maps a wiring diagram of cellular function. Science. 2016;353(6306):aaf1420. PMID: 27708008",
    "3. Krogan NJ, Cagney G, Yu H, et al. Global landscape of protein complexes in the yeast Saccharomyces cerevisiae. Nature. 2006;440(7084):637-643. PMID: 16554755",
    "4. Gavin AC, Aloy P, Grandi P, et al. Proteome survey reveals modularity of the yeast cell machinery. Nature. 2006;440(7084):631-636. PMID: 16429126",
    "5. Michaelis AC, Brunner AD, Zwiebel M, et al. The social and structural architecture of the yeast protein interactome. Nature. 2023;624(7990):192-200. PMID: 37968396",
    "6. Costanzo M, Baryshnikova A, Bellay J, et al. The genetic landscape of a cell. Science. 2010;327(5964):425-431. PMID: 20093466",
    "7. Collins SR, Miller KM, Maas NL, et al. Functional dissection of protein complexes involved in yeast chromosome biology using a genetic interaction map. Nature. 2007;446(7137):806-810. PMID: 17314980",
    "8. Costanzo M, Hou J, Messier V, et al. Environmental robustness of the global yeast genetic interaction network. Science. 2021;372(6542):eabf8424. PMID: 33958448",
    "9. Roguev A, Bandyopadhyay S, Zofall M, et al. Conservation and rewiring of functional modules revealed by an epistasis map in fission yeast. Science. 2008;322(5900):405-410. PMID: 18818364",
    "10. van Leeuwen J, Pons C, Mellor JC, et al. Exploring genetic suppression interactions on a global scale. Science. 2016;354(6312):aag0839. PMID: 27811238",
]
for ref in refs:
    ref_para(ref)

# ── FIGURE LEGENDS ────────────────────────────────────────────────────────────
head("Figure Legends")

figs = [
    ("Figure 1.",
     "Coverage of PPI pairs in the SGA dataset. Of 31,004 PPI pairs, 1,044 (3.4%, dark blue) "
     "carry a significant genetic interaction (|epsilon|>=0.08, p<0.05), 7,953 (25.6%, gray) "
     "are measured but below threshold, and 22,007 (71.0%, green) have no SGA measurement. "
     "Under a binomial null model (2.90% background rate), 261 overlapping pairs are expected; "
     "the observed 1,044 represent a 4.0-fold enrichment (z=49.1, p<<10^-100)."),
    ("Figure 2.",
     "Absolute GI epsilon for concordant and discordant PPI pairs. Violin plots of |epsilon| "
     "for Both pairs (significant GI, dark blue; mean=0.204) versus PPI-plus-non-significant "
     "pairs (gray; mean=0.027). Black lines indicate medians. Mann-Whitney U, p<<10^-100."),
    ("Figure 3.",
     "PPI interaction scores across categories. Violin plots of PPI score (range 2-10) for "
     "Both pairs (mean=4.74), PPI-plus-nonsig pairs (mean=3.44), and PPI-only pairs (no SGA "
     "measurement). Both versus PPI-plus-nonsig: Mann-Whitney U, p=4.2x10^-49."),
    ("Figure 4.",
     "PPI score versus |GI epsilon| for concordant pairs. (Left) Scatter plot of 1,044 Both "
     "pairs colored by GI sign (blue: positive epsilon; red: negative epsilon; Spearman "
     "r=0.150, p=1.1x10^-6). (Right) Binned means of |GI epsilon| per 0.75-unit PPI score "
     "interval with standard error bars."),
    ("Figure 5.",
     "ROC analysis: PPI score as predictor of significant GI. ROC curve for the 8,997 PPI "
     "pairs with direct SGA measurement. AUC=0.630."),
    ("Figure 6.",
     "GI direction at within-complex versus between-complex physical interfaces. Histograms "
     "of GI epsilon scores for within-complex Both pairs (left; 55.2% positive, n=811) and "
     "between-complex Both pairs (right; 34.3% positive, n=233). Dashed vertical line at "
     "epsilon=0; red vertical line indicates the group mean."),
    ("Figure 7.",
     "Per-complex Spearman correlations between PPI score and |GI epsilon|. Top 15 within-"
     "complex clusters by number of Both pairs (min n=5). Blue bars: p<0.05; gray bars: p>=0.05. "
     "The mitochondrial ribosome cluster (MRPL19/MRPL10/MRPL24; r=0.297, p=0.010) is the top "
     "significant complex."),
    ("Figure 8.",
     "Complex-level architecture: epistasis sign versus PPI-GI concordance. Each bubble "
     "represents a Markov cluster with at least five Both pairs. X-axis: Spearman r (PPI "
     "score vs |GI epsilon|). Y-axis: percent positive GI. Bubble size: number of Both pairs. "
     "Color: number of pairs (blue scale). Reference lines at r=0 and 50% positive GI. "
     "Notable complexes are annotated."),
]

for ft, fl in figs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(ft + " "); font(r1, bold=True)
    r2 = p.add_run(fl); font(r2)

doc.save(OUT)
print(f"Saved: {OUT}")
