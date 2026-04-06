"""
06_write_paper_v2.py  — Round 1 (Codex-role) full rewrite
Key fixes vs v1:
  1. Corrected GI threshold: |eps|>=0.08 AND p<0.05 (Costanzo 2016 standard)
  2. Fixed Fisher exact interpretation (within-cluster enriched in Both, OR=2.39)
  3. Added null model (4.0x fold-enrichment, z=49.1)
  4. Shortened abstract (~175 words)
  5. New declarative title
  6. Between-complex contacts: 65.7% negative GI -- lead finding in Results
  7. Replaced misused ref [7] with Collins 2007 (PMID 17314980) for between-pathway model
  8. Added ROC AUC=0.630 (corrected from 0.564)
  9. Improved discussion with generalizable model
  10. Fixed writing quality (passive, hedges, wordiness)
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/han-yingjhuang/Claude_projects/Genetic-interaction/paper/2026.04.06_Genetic_vs_Protein_Interactions_Yeast_v2.docx"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

def font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic

def body(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); font(r, bold=bold, italic=italic); return p

def head(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level==1 else 10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); font(r, size=14 if level==1 else 12, bold=True); return p

def center(text, size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); font(r, size=size, bold=bold, italic=italic); return p

def blank():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(""); font(r)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Physical Interaction Interfaces Concentrate Genetic Couplings, with Inter-Complex Contacts Enriched for Negative Epistasis in Yeast")
font(r, size=16, bold=True)

blank()
center("Comparative analysis of the Saccharomyces cerevisiae synthetic genetic array and proteome-scale protein interactome", size=12, italic=True)
blank(); blank()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
head("Abstract")
body(
    "Genetic interaction networks and protein-protein interaction (PPI) networks each offer a distinct view of "
    "cellular organization, yet their relationship remains poorly quantified. We systematically compared the "
    "near-complete yeast synthetic genetic array (SGA) dataset of Costanzo et al. (2016) with the proteome-scale "
    "PPI map of Michaelis et al. (2023), which covers 31,004 physical interactions among 3,927 proteins. Applying "
    "the standard SGA significance threshold (|epsilon| >= 0.08, p < 0.05), we find that 1,044 PPI pairs (3.4%) "
    "also carry a significant genetic interaction, representing a 4.0-fold enrichment over the genome-wide "
    "background rate (z = 49.1). Within-complex protein pairs are 2.4-fold more enriched at this intersection "
    "than between-complex pairs (Fisher exact, p = 2.3 x 10^-32). Strikingly, physical contacts between "
    "different protein complexes show predominantly negative genetic interactions (65.7%), whereas within-complex "
    "contacts are near-balanced (55.2% positive). Physical interaction score modestly but significantly predicts "
    "genetic interaction magnitude for within-complex pairs (Spearman r = 0.143, p = 4.6 x 10^-5) but not for "
    "between-complex pairs (r = 0.044, p = 0.50). These results reveal that the two networks are largely "
    "orthogonal but not independent, and that protein complex architecture is the primary determinant of both "
    "the likelihood and the sign of a genetic interaction at a physical interface."
)
blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(4)
r1 = p.add_run("Keywords: "); font(r1, bold=True)
r2 = p.add_run("genetic interactions, protein-protein interactions, Saccharomyces cerevisiae, epistasis, protein complexes, SGA, interactome")
font(r2)

doc.add_page_break()

# ── INTRODUCTION ─────────────────────────────────────────────────────────────
head("Introduction")

body(
    "Two large-scale technologies have mapped the functional organization of Saccharomyces cerevisiae: "
    "synthetic genetic array (SGA) analysis, which reveals functional redundancies through fitness defects "
    "in double-deletion strains [1,2], and affinity-based proteomics, which catalogs direct physical contacts "
    "between proteins [3,4,5]. At their best, these technologies provide complementary perspectives on how "
    "genes cooperate to sustain growth. Yet because they interrogate fundamentally different properties "
    "(functional buffering versus physical binding), the degree to which their networks agree or disagree, "
    "and what disagreements mean biologically, has remained unresolved."
)

body(
    "Genetic interactions (GIs) are quantified as epsilon scores: the deviation of a double mutant's fitness "
    "from the product of its single-mutant fitnesses [1]. A negative epsilon (synthetic sick or synthetic "
    "lethal) indicates that the two genes buffer each other's loss, typically because they act in parallel "
    "pathways or compensatory mechanisms. A positive epsilon indicates that the two genes function in the "
    "same pathway, rendering the second mutation phenotypically inconsequential. Costanzo et al. (2010) and "
    "Costanzo et al. (2016) systematically measured GIs across the yeast non-essential gene space, generating "
    "approximately 23 million double-mutant fitness measurements and revealing that functionally related "
    "genes cluster by GI profile, with protein complexes forming dense intra-cluster modules [2,6]."
)

body(
    "Protein-protein interaction (PPI) maps have undergone parallel expansions. Krogan et al. (2006) and "
    "Gavin et al. (2006) used tandem affinity purification coupled with mass spectrometry to resolve the "
    "yeast protein complex landscape [3,4]. Michaelis et al. (2023) extended this work by applying "
    "high-throughput affinity enrichment mass spectrometry across 4,159 GFP-tagged strains, producing a "
    "near-complete interactome of 31,004 interactions among 3,927 proteins with a composite confidence "
    "score derived from three independent evidence layers [5]. This dataset, organized into 617 Markov "
    "clusters, provides an opportunity to ask how physical interaction strength maps onto genetic coupling "
    "at both the global and per-complex levels."
)

body(
    "The conceptual relationship between GIs and PPIs has been framed around a between-pathway model: "
    "protein pairs within the same complex tend to show positive GIs because loss of any subunit impairs "
    "the same functional unit, whereas pairs in parallel pathways show negative GIs because each buffers "
    "the other's absence [2,7]. Direct quantitative tests of this model using PPI scores as a continuous "
    "variable, rather than binary complex co-membership, have not been performed with the most comprehensive "
    "datasets now available."
)

body(
    "Here we carry out a systematic, genome-scale comparison of the Costanzo et al. (2016) SGA network and "
    "the Michaelis et al. (2023) PPI network. We address whether PPI interfaces are enriched for genetic "
    "interactions relative to a null model, whether discordant pairs carry residual genetic signals, how "
    "physical interaction strength correlates with genetic interaction magnitude, and how complex membership "
    "governs the sign of GIs at physical interfaces. The answers provide a quantitative foundation for "
    "integrating the two networks and illuminate a layered architecture in which physical contact is "
    "necessary but not sufficient for detectable genetic coupling."
)

# ── RESULTS ──────────────────────────────────────────────────────────────────
head("Results")

head("Protein-protein interaction interfaces are 4-fold enriched for genetic interactions", level=2)

body(
    "The Michaelis et al. (2023) dataset contains 31,004 physical interactions among 3,927 S. cerevisiae "
    "proteins organized into 617 Markov clusters. The Costanzo et al. (2016) SGA dataset provides "
    "12,698,939 gene-pair measurements across the non-essential gene space. Of these measurements, "
    "10,707,139 (84.3%) involve at least one gene whose protein product appears in the PPI network, "
    "yielding 7,721,253 unique tested pairs. Applying the standard significance threshold of "
    "|epsilon| >= 0.08 and p < 0.05, 223,155 pairs carry a significant genetic interaction, giving a "
    "background rate of 2.90%."
)

body(
    "Of the 31,004 PPI pairs, 8,997 are directly tested by SGA (Figure 1). Among these, 1,044 (3.4%) "
    "carry a significant genetic interaction, whereas 7,953 (25.6%) are measured but fall below the "
    "significance threshold. Under a binomial null model with the 2.90% background rate, only 261 "
    "overlapping pairs would be expected by chance, making the observed 1,044 a 4.0-fold enrichment "
    "(z = 49.1, p << 10^-100). The remaining 22,007 PPI pairs (71.0%) have no SGA counterpart, reflecting "
    "the restriction of the current SGA dataset to non-essential-by-non-essential gene pairs. In the "
    "other direction, 223,155 significant GI pairs lack a corresponding PPI, confirming that most genetic "
    "buffering operates between non-contacting genes."
)

head("Discordant pairs carry sub-threshold rather than absent genetic coupling", level=2)

body(
    "We asked whether PPI pairs lacking a significant GI are biologically uncoupled or merely weakly "
    "coupled. The mean absolute epsilon for PPI pairs with measured but non-significant GIs is 0.027, "
    "compared with 0.204 for pairs with both a PPI and significant GI (Mann-Whitney U, p << 10^-100; "
    "Figure 2). This 7.5-fold difference indicates that discordant PPI pairs retain measurable but "
    "sub-threshold genetic signal. Consistently, PPI scores for concordant pairs are also higher "
    "(mean 4.74 versus 3.44; Mann-Whitney U, p = 4.2 x 10^-49; Figure 3), indicating that "
    "higher-confidence physical interactions are more likely to be accompanied by detectable genetic coupling."
)

body(
    "Despite these trends, PPI score alone is an imperfect predictor of GI significance. A receiver "
    "operating characteristic analysis using PPI score to classify pairs as significant-GI versus "
    "non-significant-GI yields an AUC of 0.630 (Figure 5), indicating that physical interaction "
    "strength captures only part of the variance in genetic coupling. The majority of PPI pairs remain "
    "below the GI significance threshold regardless of their confidence score, underscoring that "
    "additional factors, most likely pathway context and the availability of compensatory routes, "
    "determine whether a physical contact gives rise to a detectable genetic interaction."
)

head("Inter-complex physical contacts are enriched for negative epistasis", level=2)

body(
    "Among the 1,044 concordant pairs (Both), we stratified by whether the two interacting proteins "
    "belong to the same Markov cluster (within-complex, n = 811) or different clusters "
    "(between-complex, n = 233). Within-complex pairs account for 77.7% of Both pairs, compared with "
    "59.3% of PPI-plus-nonsig pairs, a significant enrichment (Fisher exact, OR = 2.39, p = 2.3 x 10^-32). "
    "Within-complex contacts are therefore 2.4-fold more likely to generate a detectable genetic "
    "interaction than between-complex contacts."
)

body(
    "The GI direction at these two classes of interface diverges markedly (Figure 6). Within-complex "
    "Both pairs are 55.2% positive, consistent with the expectation that same-complex partners function "
    "in the same pathway. Between-complex Both pairs, by contrast, are only 34.3% positive (65.7% "
    "negative; Figure 6). Physical contacts that bridge different functional modules are thus enriched "
    "for synthetic sick or synthetic lethal relationships, suggesting that inter-complex bridges couple "
    "two functional units such that disrupting either partner impairs both simultaneously. This pattern "
    "mirrors the between-pathway negative epistasis model proposed from E-MAP studies [7] but now "
    "identifies it specifically at the subset of inter-complex contacts detectable by co-purification."
)

head("Physical interaction strength predicts genetic interaction magnitude within but not between complexes", level=2)

body(
    "For all 1,044 Both pairs, PPI score and absolute GI epsilon show a modest but significant positive "
    "correlation (Spearman r = 0.150, p = 1.1 x 10^-6; Pearson r = 0.140, p = 5.8 x 10^-6; Figure 4). "
    "The correlation is stronger for positive GIs (Spearman r = 0.258, p = 1.7 x 10^-9, n = 528) than "
    "for negative GIs (r = 0.185, p = 2.5 x 10^-5, n = 516), consistent with the idea that stronger "
    "physical coupling more faithfully predicts stronger same-pathway functional dependence."
)

body(
    "Stratifying by complex membership reveals a critical asymmetry. For within-complex pairs, PPI "
    "score significantly predicts GI magnitude (Spearman r = 0.143, p = 4.6 x 10^-5; mean |epsilon| = "
    "0.212). For between-complex pairs, no significant correlation is observed (r = 0.044, p = 0.50; "
    "mean |epsilon| = 0.176). Thus the relationship between physical interaction strength and genetic "
    "interaction magnitude is confined to intra-complex contacts, where PPI score reflects structural "
    "or functional indispensability of the contact. Between-complex physical contacts that generate "
    "strong GIs do so through a mechanism that is decoupled from interaction confidence."
)

head("Complex identity determines epistasis pattern at physical interfaces", level=2)

body(
    "Grouping within-complex Both pairs by Markov cluster revealed 214 clusters with at least one "
    "concordant pair, of which 68 contained five or more pairs suitable for cluster-level analysis "
    "(Figure 7). Per-cluster Spearman correlations between PPI score and |GI epsilon| ranged from "
    "r = -0.35 to r = 0.30, with the mitochondrial ribosome cluster (MRPL19/MRPL10/MRPL24; n = 75) "
    "showing the strongest significant positive correlation (r = 0.297, p = 0.010). In contrast, "
    "the large and small cytoplasmic ribosomal subunit clusters (n = 91 and n = 99, respectively) "
    "show no significant correlation (r = 0.096, r = 0.115), consistent with the idea that all "
    "cytoplasmic ribosomal subunit contacts are equally critical, flattening the relationship "
    "between PPI score and GI magnitude."
)

body(
    "The sign of GIs within complexes also varies systematically. The spliceosome cluster "
    "(PRP4/DIB1/SNU66; n = 43) shows 91% negative GI, indicating that spliceosome subunit pairs "
    "are predominantly synthetic sick or synthetic lethal, consistent with the near-essential nature "
    "of mRNA splicing. The nuclear pore complex cluster (POM152/MTR2/NUP188; n = 14) shows 100% "
    "negative GI. In contrast, the Sin3-associated chromatin regulatory cluster (STB3/CTI6/SIN3; "
    "n = 24) shows 96% positive GI and the elongator complex (ELP2/ELP3/ELP4; n = 16) shows "
    "100% positive GI, consistent with these modules functioning as integrated units in which "
    "any component loss produces the same functional deficit (Figure 8). These observations suggest "
    "that the balance between positive and negative GI within a complex reflects the essentiality "
    "of the complex's function and the degree of functional redundancy among its components."
)

# ── DISCUSSION ───────────────────────────────────────────────────────────────
head("Discussion")

body(
    "This analysis establishes that protein-protein interaction interfaces are significantly enriched "
    "for genetic interactions (4.0-fold over background), but that this enrichment is asymmetric: "
    "within-complex contacts are 2.4 times more enriched than between-complex contacts, and the two "
    "classes of contact generate qualitatively different epistasis patterns. These results support a "
    "layered model in which physical contact is a necessary predisposing factor for genetic coupling, "
    "but complex architecture and pathway context determine whether the coupling is detectable and "
    "whether it is positive or negative."
)

body(
    "The modest global correlation between PPI score and GI magnitude (r = 0.150) reflects genuine "
    "biological complexity rather than measurement noise. The correlation is driven almost entirely "
    "by within-complex pairs; for between-complex contacts, PPI score does not predict GI magnitude. "
    "This dissociation suggests that the two quantities measure distinct properties at inter-complex "
    "interfaces: PPI score reports on co-purification efficiency or abundance correlation, neither of "
    "which necessarily tracks with the severity of combined pathway disruption. Future interaction "
    "scoring systems might benefit from incorporating genetic interaction context to better capture "
    "the functional weight of a physical contact."
)

body(
    "The enrichment of negative GIs at between-complex physical interfaces (65.7%) is particularly "
    "informative. These contacts, which represent proteins that physically bridge two distinct "
    "functional modules, tend to be synthetic sick when both components are disrupted. One "
    "interpretation is that inter-complex bridging proteins serve as functional conduits between "
    "modules, so that losing either the bridge protein or a component of the module it connects "
    "can impair the same downstream process. Alternatively, inter-complex contacts may identify "
    "non-stoichiometric regulatory contacts where the interacting proteins independently regulate "
    "a shared target, generating negative GIs through parallel pathway logic. Distinguishing "
    "between these mechanisms will require further characterization of the specific inter-complex "
    "contacts identified here."
)

body(
    "The near-complete orthogonality of the two networks (only 3.4% of PPI pairs carry a significant "
    "GI) has practical implications. It argues against using one network as a simple proxy for the "
    "other and in favor of treating them as independent evidence layers in multi-modal network models. "
    "The 4.0-fold enrichment nonetheless makes PPI data a useful filter for prioritizing candidate "
    "gene pairs in targeted GI screens, particularly within well-characterized complexes such as the "
    "mitochondrial ribosome, where PPI score is genuinely predictive of GI magnitude. Conversely, "
    "the 223,155 significant GI pairs without any PPI counterpart represent a large space of "
    "functional relationships that are invisible to co-purification approaches and that can only "
    "be explored through genetic perturbation."
)

body(
    "The complex-specific epistasis patterns add a new layer of resolution to the between-pathway model "
    "of genetic interactions. It was previously known that proteins in the same complex tend toward "
    "positive GIs and that proteins in parallel pathways tend toward negative GIs [2,7]. The present "
    "data reveal that this rule does not apply uniformly even within complexes: the spliceosome and "
    "nuclear pore complex are predominantly synthetic sick, whereas the elongator complex and the "
    "Sin3-associated chromatin complex are predominantly positive. These differences likely reflect "
    "the degree to which individual components are rate-limiting for complex function. In complexes "
    "where all subunits are equally essential for activity (such as the spliceosome), losing any two "
    "creates a synthetic phenotype. In complexes where subunit loss causes a partial, graded impairment "
    "(such as elongator), the double mutant is not more severe than either single mutant."
)

body(
    "Several limitations constrain the current analysis. The SGA dataset covers only non-essential "
    "gene pairs; essential genes, which constitute roughly 20% of the yeast genome and are "
    "disproportionately present in the PPI network, are not represented. The 71% of PPI pairs with "
    "no SGA measurement are enriched for essential-essential and essential-non-essential pairs, "
    "so our estimates of overlap and enrichment apply primarily to the non-essential proteome. "
    "Including data from temperature-sensitive and DAmP allele screens of essential genes [6] "
    "would substantially increase coverage and may reveal higher overlap fractions. Additionally, "
    "GI profiles are condition-dependent [8], and integration with condition-specific GI data "
    "may reveal cryptic concordances obscured by averaging across standard growth conditions. "
    "Finally, the PPI scoring system used here reflects co-purification efficiency, which is "
    "a proxy for binding affinity but not a direct measure of it; direct structural data on "
    "binding energetics would allow a more mechanistic interpretation of the correlation with GI magnitude."
)

# ── METHODS ──────────────────────────────────────────────────────────────────
head("Methods")

head("Data acquisition", level=2)
body(
    "Genetic interaction data were obtained from Costanzo et al. (2016) [2] via the CellMap repository "
    "(thecellmap.org). The file SGA_NxN.txt (12,698,939 rows) contains pairwise genetic interaction "
    "data for non-essential S. cerevisiae genes, including query and array strain identifiers, "
    "epsilon score, p-value, single-mutant fitnesses, and double-mutant fitness. Protein-protein "
    "interaction data were taken from Supplementary Data 2 of Michaelis et al. (2023) [5], "
    "comprising a semicolon-delimited edge file (31,004 edges) and a node file (3,927 proteins), "
    "both using Saccharomyces Genome Database (SGD) ORF identifiers as primary keys. Markov "
    "cluster assignments from the node file (617 clusters) were used as the definition of "
    "protein complex membership throughout."
)

head("Data preprocessing and matching", level=2)
body(
    "SGA strain identifiers (format: YXXXXX_sn123) were parsed to extract ORF identifiers using "
    "a regular expression matching the standard SGD ORF pattern. SGA data were processed in "
    "chunks of 200,000 rows; only rows in which at least one partner ORF appeared in the PPI "
    "network were retained (10,707,139 of 12,698,939 rows). Multiple SGA measurements for the "
    "same gene pair were aggregated by computing the mean epsilon score and the minimum p-value "
    "across replicates. Gene pairs and PPI edges were represented as frozensets of sorted ORF "
    "identifiers to ensure direction-independent matching. A genetic interaction was classified "
    "as significant if both |epsilon| >= 0.08 and p < 0.05, following the standard threshold "
    "used in Costanzo et al. (2016) [2]."
)

head("Statistical analyses", level=2)
body(
    "The null model for the expected number of overlapping pairs was computed as a binomial "
    "expectation: background_rate x n_measured_PPI_pairs, where the background rate is the "
    "fraction of SGA-measured pairs that pass the significance threshold (223,155 / 7,721,253 "
    "= 2.90%). The observed-versus-expected enrichment was tested using a normal approximation "
    "to the binomial. Differences in |epsilon| and PPI score across categories were assessed "
    "using the two-sided Mann-Whitney U test. Within-cluster versus between-cluster enrichment "
    "in the Both category was tested using Fisher's exact test, with the contingency table "
    "defined by category (Both vs PPI-plus-nonsig) and cluster status (same vs different "
    "Markov cluster). Spearman and Pearson correlations were computed using scipy.stats. "
    "Per-cluster correlations were computed for clusters with at least five Both pairs. "
    "The ROC AUC was computed using scikit-learn, treating significant GI status as the "
    "positive class and PPI score as the continuous predictor, in the subset of PPI pairs "
    "with a direct SGA measurement."
)

head("Figures", level=2)
body(
    "All figures were generated using matplotlib and seaborn (Python). Violin plots show "
    "kernel density estimates with medians. Scatter plots use jitter on the PPI score axis "
    "for visual clarity. Correlation bar charts display Spearman r values colored by "
    "significance (p < 0.05 versus p >= 0.05). The complex-level summary plot (Figure 8) "
    "encodes per-cluster Spearman r on the x-axis and the fraction of positive GIs on the "
    "y-axis, with bubble size proportional to the number of concordant pairs. All figures "
    "were saved at 300 DPI."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
head("References")

refs = [
    "1. Tong AH, Evangelista M, Parsons AB, et al. Systematic genetic analysis with ordered arrays of yeast deletion mutants. Science. 2001;294(5550):2364-2368. PMID: 11743205",
    "2. Costanzo M, VanderSluis B, Koch EN, et al. A global genetic interaction network maps a wiring diagram of cellular function. Science. 2016;353(6306):aaf1420. PMID: 27708008",
    "3. Krogan NJ, Cagney G, Yu H, et al. Global landscape of protein complexes in the yeast Saccharomyces cerevisiae. Nature. 2006;440(7084):637-643. PMID: 16554755",
    "4. Gavin AC, Aloy P, Grandi P, et al. Proteome survey reveals modularity of the yeast cell machinery. Nature. 2006;440(7084):631-636. PMID: 16429126",
    "5. Michaelis AC, Brunner AD, Zwiebel M, et al. The social and structural architecture of the yeast protein interactome. Nature. 2023;624(7990):192-200. PMID: 37968396",
    "6. Costanzo M, Baryshnikova A, Bellay J, et al. The genetic landscape of a cell. Science. 2010;327(5964):425-431. PMID: 20093466",
    "7. Collins SR, Miller KM, Maas NL, et al. Functional dissection of protein complexes involved in yeast chromosome biology using a genetic interaction map. Nature. 2007;446(7137):806-810. PMID: 17314980",
    "8. Costanzo M, Hou J, Messier V, et al. Environmental robustness of the global yeast genetic interaction network. Science. 2021;372(6542):eabf8424. PMID: 33958448",
    "9. Roguev A, Bandyopadhyay S, Zofall M, et al. Conservation and rewiring of functional modules revealed by an epistasis map in fission yeast. Science. 2008;322(5900):405-410. PMID: 18818364",
    "10. van Leeuwen J, Pons C, Mellor JC, et al. Exploring genetic suppression interactions on a global scale. Science. 2016;354(6312):aag0839. PMID: 27811238",
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    r = p.add_run(ref); font(r, size=10)

# ── FIGURE LEGENDS ────────────────────────────────────────────────────────────
head("Figure Legends")

fig_legs = [
    ("Figure 1.", "Coverage of PPI pairs in the SGA dataset. Pie chart showing 31,004 PPI pairs classified as Both (PPI and significant GI, |epsilon|>=0.08, p<0.05; dark blue, 3.4%), PPI with measured but non-significant GI (gray, 25.6%), and PPI with no SGA measurement (green, 71.0%). Dashed annotation indicates 4.0-fold enrichment over the 2.90% background GI rate (z=49.1)."),
    ("Figure 2.", "Absolute GI epsilon scores for concordant and discordant PPI pairs. Violin plots comparing |epsilon| for Both pairs (significant GI, dark blue) versus PPI-plus-non-significant GI pairs (gray). Black lines indicate medians. Both-pair mean = 0.204, non-sig mean = 0.027 (Mann-Whitney U, p << 10^-100)."),
    ("Figure 3.", "PPI interaction scores across categories. Violin plots showing the distribution of PPI scores (range 2-10) for Both pairs (mean=4.74), PPI-plus-nonsig pairs (mean=3.44), and PPI-only (no SGA measurement) pairs. Black lines indicate medians."),
    ("Figure 4.", "PPI score versus |GI epsilon| for concordant pairs. Left: scatter plot of 1,044 Both pairs colored by GI sign (blue, positive; red, negative; Spearman r=0.150, p=1.1x10^-6). Right: binned means of |GI epsilon| per PPI score interval with standard error bars."),
    ("Figure 5.", "ROC analysis: PPI score as predictor of significant GI status. ROC curve computed in the 8,997 PPI pairs with direct SGA measurement. AUC=0.630."),
    ("Figure 6.", "GI direction at within-complex versus between-complex physical interfaces. Histograms of GI epsilon scores for within-complex Both pairs (left, 55.2% positive) and between-complex Both pairs (right, 34.3% positive). Dashed vertical line at epsilon=0; red line indicates mean."),
    ("Figure 7.", "Per-complex Spearman correlations between PPI score and |GI epsilon|. Top 15 within-complex cluster groups by number of Both pairs (min n=5), colored by significance (blue, p<0.05; gray, p>=0.05). MRPL cluster (r=0.297, p=0.010) is the top significant complex."),
    ("Figure 8.", "Complex-level architecture of epistasis at physical interfaces. Each bubble represents a Markov cluster with at least five Both pairs; x-axis, Spearman r (PPI score vs |GI epsilon|); y-axis, percent positive GI; bubble size proportional to number of Both pairs. Color encodes number of pairs (blue scale). Quadrant lines at r=0 and 50% positive GI. Notable complexes are annotated."),
]

for ft, fl in fig_legs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(ft + " "); font(r1, bold=True)
    r2 = p.add_run(fl); font(r2)

doc.save(OUT)
print(f"Saved: {OUT}")
