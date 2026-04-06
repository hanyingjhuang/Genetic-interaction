"""
08_write_paper_v4.py — Final version (joint Codex+Claude polish pass)
Key fixes vs v3:
  1. Corrected between-complex GI sign: 34.3% positive ≈ background 35.6% (OR=0.95, p=0.73, NOT enriched)
     The real finding: within-complex contacts are enriched 1.55x for positive GI (OR=2.23, p=7e-30)
  2. Fixed Discussion para 3 accordingly (removed false between-complex negative enrichment claim)
  3. Testable prediction revised for the corrected finding
  4. Discussion opener rewritten to lead with the finding, not self-reference
  5. Removed imprecise "1 in 3 contacts fail" statement; replaced with AUC-based framing
  6. Added citation for genome-wide GI sign fractions (Costanzo 2016)
  7. Minor prose polish: passive voice and hedges removed
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/han-yingjhuang/Claude_projects/Genetic-interaction/paper/2026.04.06_Genetic_vs_Protein_Interactions_Yeast_v4.docx"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

def font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic

def body(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); font(r, bold=bold, italic=italic); return p

def head(text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); font(r, size=14 if level == 1 else 12, bold=True); return p

def center(text, size=11, bold=False, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); font(r, size=size, bold=bold, italic=italic); return p

def blank():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.add_run("").font.name = "Times New Roman"

def ref_para(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    r = p.add_run(text); font(r, size=10)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
blank()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run(
    "Protein Complex Membership Governs the Sign and Detectability of Epistasis "
    "at Physical Interaction Interfaces in Yeast"
); font(r, size=16, bold=True)
blank()
center(
    "A genome-scale comparison of the Saccharomyces cerevisiae synthetic genetic array "
    "and proteome-scale protein interactome",
    size=12, italic=True
)
blank(); blank()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
head("Abstract")
body(
    "Genetic interactions and protein-protein interactions (PPIs) each reflect how genes and "
    "proteins cooperate to sustain cellular function, yet their quantitative relationship remains "
    "poorly defined. We compared the near-complete yeast synthetic genetic array (SGA) dataset of "
    "Costanzo et al. (2016) with the proteome-scale PPI map of Michaelis et al. (2023), which "
    "covers 31,004 physical interactions among 3,927 proteins. Using the standard SGA significance "
    "threshold (|epsilon| >= 0.08, p < 0.05), PPI interfaces concentrate genetic interactions at "
    "4.0-fold the genome-wide background rate (1,044 of 8,997 tested pairs; z = 49.1). Within-complex "
    "protein pairs are 2.4-fold more enriched at this intersection than between-complex pairs "
    "(Fisher exact, OR = 2.39, p = 2.3 x 10^-32). Within-complex physical contacts are strongly "
    "enriched for positive epistasis (55.2% positive versus 35.6% genome-wide background; OR = 2.23, "
    "p = 7.0 x 10^-30), whereas between-complex contacts show no departure from the background sign "
    "distribution (34.3% positive, OR = 0.95, p = 0.73). Physical interaction score significantly "
    "predicts genetic interaction magnitude within complexes (Spearman r = 0.143, p = 4.6 x 10^-5) "
    "but not between complexes (r = 0.044, p = 0.50). These findings establish that protein complex "
    "membership, rather than the fact of physical association, is the primary determinant of both "
    "the likelihood and the sign of epistasis at physical interaction interfaces."
)
blank()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(4)
r1 = p.add_run("Keywords: "); font(r1, bold=True)
r2 = p.add_run(
    "genetic interactions, protein-protein interactions, Saccharomyces cerevisiae, "
    "epistasis, protein complexes, SGA, interactome"
); font(r2)

doc.add_page_break()

# ── INTRODUCTION ─────────────────────────────────────────────────────────────
head("Introduction")

body(
    "Two large-scale technologies have mapped the functional organization of Saccharomyces cerevisiae. "
    "Synthetic genetic array (SGA) analysis detects functional redundancies by measuring fitness "
    "defects in double-deletion strains [1,2]. Affinity-based proteomics catalogs direct physical "
    "contacts between proteins [3,4,5]. Together, these approaches have generated the most "
    "comprehensive functional genomics datasets available for any eukaryote. Because they interrogate "
    "fundamentally different properties, the degree to which their networks agree or diverge, and what "
    "disagreements imply biologically, has remained largely unresolved."
)

body(
    "SGA quantifies genetic interactions (GIs) as epsilon scores: the deviation of a double mutant's "
    "fitness from the product of its single-mutant fitnesses [1]. Negative epsilon values "
    "(synthetic sick or lethal) indicate that the two genes buffer each other's loss, typically through "
    "parallel or redundant mechanisms. Positive values indicate that both genes impair the same functional "
    "pathway, rendering the second mutation phenotypically inconsequential. Costanzo et al. (2010) and "
    "Costanzo et al. (2016) applied this approach at genome scale, generating approximately 23 million "
    "double-mutant measurements across the non-essential gene space and revealing that functionally "
    "related genes cluster by GI profile, with protein complexes occupying the cores of dense "
    "intra-cluster modules [2,6]."
)

body(
    "Protein-protein interaction maps have undergone parallel expansions. Krogan et al. (2006) and "
    "Gavin et al. (2006) used tandem affinity purification coupled with mass spectrometry to resolve "
    "the yeast protein complex landscape [3,4]. Michaelis et al. (2023) extended this to a near-complete "
    "interactome of 31,004 interactions among 3,927 proteins, using three independent evidence layers "
    "to produce a composite confidence score and organizing the network into 617 Markov clusters that "
    "define functional complex assignments [5]. This dataset offers the first opportunity to test, at "
    "near-complete resolution, how physical interaction strength maps onto genetic coupling at both "
    "global and per-complex levels."
)

body(
    "Conceptually, the between-pathway model predicts that proteins within the same complex show "
    "positive GIs, because any subunit loss impairs the same functional unit, whereas proteins in "
    "parallel pathways show negative GIs through mutual functional buffering [2,7]. Whether this "
    "model applies uniformly to the subset of between-pathway pairs that make physical contact, and "
    "whether physical interaction confidence predicts GI magnitude within and between complexes, are "
    "questions that prior work has not addressed with continuous, proteome-scale data. Here we perform "
    "this comparison directly and find that protein complex membership, not the fact of physical "
    "contact, determines both the likelihood and the sign of epistasis at a physical interface."
)

# ── RESULTS ──────────────────────────────────────────────────────────────────
head("Results")

head("PPI interfaces concentrate genetic interactions at 4-fold the genome-wide background rate", level=2)

body(
    "The Michaelis et al. (2023) dataset contains 31,004 physical interactions among 3,927 "
    "S. cerevisiae proteins organized into 617 Markov clusters. Of 12,698,939 SGA gene-pair "
    "measurements, 10,707,139 involve at least one gene product present in the PPI network, "
    "yielding 7,721,253 unique tested pairs. Using the standard significance threshold of "
    "|epsilon| >= 0.08 and p < 0.05 [2], 223,155 pairs carry a significant GI (background "
    "rate: 2.90%)."
)

body(
    "Of the 31,004 PPI pairs, 8,997 are directly tested by SGA: 1,044 carry a significant GI "
    "(3.4%) and 7,953 are measured but below threshold (Figure 1). Under a binomial null model "
    "at the 2.90% background rate, only 261 overlapping pairs are expected by chance. The "
    "observed 1,044 represent a 4.0-fold enrichment (z = 49.1, p < 10^-100). The remaining "
    "22,007 PPI pairs (71.0%) lack SGA measurements, predominantly because one or both partners "
    "are essential genes excluded from the non-essential SGA screen. In the other direction, "
    "223,155 significant GI pairs lack a PPI counterpart, confirming that most genetic buffering "
    "operates between non-contacting proteins."
)

head("Discordant PPI pairs carry sub-threshold genetic signal", level=2)

body(
    "Among PPI pairs with direct SGA measurements, those lacking a significant GI are not "
    "biologically silent. The mean absolute epsilon for PPI-measured but non-significant pairs "
    "is 0.027, compared with 0.204 for concordant (Both) pairs (Mann-Whitney U, p < 10^-100; "
    "Figure 2). This 7.5-fold difference establishes that discordant pairs retain residual "
    "genetic signal below the detection threshold. Concordant pairs also receive higher PPI "
    "interaction scores (mean 4.74 versus 3.44; Mann-Whitney U, p = 4.2 x 10^-49; Figure 3), "
    "indicating that higher-confidence physical interactions are modestly more likely to produce "
    "detectable genetic coupling."
)

body(
    "A receiver operating characteristic analysis using PPI score to predict GI significance "
    "in the 8,997 measured pairs yields an AUC of 0.630 (Figure 5). Although significantly "
    "above chance, this value establishes that PPI confidence alone explains only a modest "
    "fraction of the variance in genetic coupling. Pathway architecture and the availability "
    "of compensatory routes are likely the predominant determinants of whether a physical "
    "contact generates a detectable GI."
)

head("Within-complex contacts are enriched for positive epistasis; between-complex contacts are not", level=2)

body(
    "Among the 1,044 concordant pairs, 811 involve proteins from the same Markov cluster "
    "(within-complex, 77.7%) and 233 involve proteins from different clusters (between-complex, "
    "22.3%). Within-complex pairs are 2.4-fold more enriched among concordant than among "
    "non-significant pairs (Fisher exact, OR = 2.39, p = 2.3 x 10^-32), establishing that "
    "intra-complex physical contacts are substantially more likely to generate detectable "
    "genetic interactions."
)

body(
    "GI sign distributions for the two interface classes diverge in a functionally informative "
    "way (Figure 6). Within-complex Both pairs are 55.2% positive, a 1.55-fold enrichment over "
    "the genome-wide positive-GI rate of 35.6% among all significant SGA pairs (Fisher exact, "
    "OR = 2.23, p = 7.0 x 10^-30; [2]). Within-complex physical contacts thus strongly "
    "recapitulate the same-pathway positive epistasis signature. Between-complex Both pairs, "
    "by contrast, are 34.3% positive, statistically indistinguishable from the genome-wide "
    "background (OR = 0.95, p = 0.73). Physical contacts that bridge different functional "
    "modules therefore generate GIs at 4-fold the background rate, but with the same sign "
    "distribution as non-interacting gene pairs. These inter-complex bridging contacts appear "
    "to participate in between-pathway genetic buffering without any additional sign bias "
    "attributable to the physical contact itself."
)

head("Physical interaction strength predicts GI magnitude within but not between complexes", level=2)

body(
    "Across all 1,044 concordant pairs, PPI score and absolute GI epsilon show a modest but "
    "significant positive correlation (Spearman r = 0.150, p = 1.1 x 10^-6; Pearson r = 0.140, "
    "p = 5.8 x 10^-6; Figure 4). The correlation is stronger for positive GIs (Spearman r = "
    "0.258, p = 1.7 x 10^-9, n = 528) than for negative GIs (r = 0.185, p = 2.5 x 10^-5, "
    "n = 516)."
)

body(
    "Stratifying by complex membership reveals a critical asymmetry. Within-complex pairs show "
    "a significant correlation (Spearman r = 0.143, p = 4.6 x 10^-5, n = 811; mean |epsilon| = "
    "0.212), whereas between-complex pairs do not (r = 0.044, p = 0.50, n = 233; mean "
    "|epsilon| = 0.176). Physical interaction confidence therefore tracks genetic interaction "
    "magnitude only within functional modules, where co-purification efficiency plausibly "
    "reflects structural indispensability. At inter-complex interfaces, the determinants of "
    "GI magnitude are decoupled from interaction confidence."
)

head("Complex identity shapes the epistasis landscape at physical interfaces", level=2)

body(
    "Grouping within-complex concordant pairs by Markov cluster identified 214 clusters with "
    "at least one Both pair, of which 68 contained five or more pairs for per-cluster analysis "
    "(Figures 7, 8). Per-cluster Spearman correlations between PPI score and |GI epsilon| vary "
    "widely. The mitochondrial ribosome (MRPL19/MRPL10/MRPL24; n = 75) shows the strongest "
    "significant positive correlation (r = 0.297, p = 0.010), suggesting that mitochondrial "
    "ribosomal contacts vary in functional indispensability in proportion to co-purification "
    "efficiency. The large and small cytoplasmic ribosomal subunit clusters (n = 91 and 99) show "
    "no significant correlation (r = 0.096 and r = 0.115), consistent with all cytoplasmic "
    "ribosomal contacts being uniformly indispensable."
)

body(
    "The fraction of positive GIs within each complex also varies systematically with complex "
    "function. The spliceosome (PRP4/DIB1/SNU66; n = 43) shows 91% negative GI, and the nuclear "
    "pore complex (POM152/MTR2/NUP188; n = 14) shows 100% negative GI. We note that both "
    "complexes contain near-essential components, so some of this pattern may reflect the "
    "near-essential gene context rather than a distinctive property of the physical contacts "
    "themselves; further analysis with essential-gene GI data will be needed to fully "
    "disentangle these contributions. In contrast, the Sin3-associated chromatin regulatory "
    "cluster (STB3/CTI6/SIN3; n = 24) shows 96% positive GI and the elongator complex "
    "(ELP2/ELP3/ELP4; n = 16) shows 100% positive GI, consistent with these modules "
    "functioning as integrated units in which any component loss produces the same partial "
    "functional deficit. This pattern implies that the balance of positive and negative GI "
    "within a complex primarily reflects whether individual components are rate-limiting for "
    "complex activity, not whether they physically contact each other."
)

# ── DISCUSSION ───────────────────────────────────────────────────────────────
head("Discussion")

body(
    "Physical interaction interfaces in yeast are 4-fold enriched for genetic interactions, "
    "but this enrichment is not uniform: within-complex contacts are 2.4-fold more enriched "
    "than between-complex contacts, and only within-complex contacts shift the balance toward "
    "positive epistasis. Between-complex contacts generate GIs at 4-fold the background rate "
    "but with the same sign distribution as the rest of the GI network, revealing that "
    "physical contact at an inter-complex interface is neither necessary nor sufficient to "
    "bias epistasis sign. Protein complex membership, not the existence of a physical contact, "
    "governs the direction of genetic coupling."
)

body(
    "The within-complex enrichment for positive GI (1.55x; p = 7.0 x 10^-30) directly "
    "quantifies, for the first time at near-proteome scale, the same-pathway epistasis logic "
    "central to network biology: proteins that physically associate within a complex tend to "
    "function in the same pathway, so disrupting both produces no additional phenotypic cost "
    "beyond either single mutant. That this enrichment is specific to within-complex contacts "
    "and absent from between-complex contacts indicates that the spatial context of a physical "
    "interface, whether within one module or bridging two, is as informative as the interaction "
    "itself. Future interaction scoring systems that incorporate complex membership alongside "
    "co-purification confidence may therefore better capture the functional context of a "
    "physical contact."
)

body(
    "The behavior of between-complex contacts is equally informative by its neutrality. These "
    "contacts generate GIs at above-background rates yet show no bias in epistasis sign, "
    "implying that the physical bridge between two modules does not, on its own, impose a "
    "same-pathway or between-pathway logic. A plausible interpretation is that inter-complex "
    "bridging contacts are functionally heterogeneous: some connect proteins in truly parallel "
    "pathways (generating negative GIs), others connect proteins with shared outputs "
    "(generating positive GIs), and the two cancel at the population level. A testable "
    "prediction of this model is that inter-complex bridging proteins with high structural "
    "conservation across species will show more biased GI sign distributions than lineage-"
    "specific bridging contacts, because structural conservation would select for functionally "
    "homogeneous bridge types. This prediction can be tested by integrating cross-species "
    "epistasis maps [9] with the current interactome data."
)

body(
    "The complex-specific heterogeneity in epistasis architecture further refines the "
    "between-pathway model of GIs [7]. It was previously known that complex membership "
    "predicts GI sign at the population level; the present data reveal that this rule has "
    "substantial within-complex variance controlled by the essentiality of complex function "
    "and the degree of individual component dispensability. Complexes involved in near-essential "
    "processes, such as splicing and nuclear transport, are enriched for negative GIs because "
    "any pair of subunit mutations together collapses an essential function. Complexes that "
    "tolerate partial loss of individual components, such as elongator and the Sin3 complex, "
    "are enriched for positive GIs because any component loss produces the same partial "
    "impairment. This distinction, between essential-complex negative epistasis and "
    "dispensable-complex positive epistasis, is likely to generalize to other organisms given "
    "the conservation of epistasis patterns observed in fission yeast [9]."
)

body(
    "The modest PPI-to-GI correlation (r = 0.150 globally; AUC = 0.630) has clear "
    "implications for multi-omic network integration. PPI data can prioritize candidate gene "
    "pairs for targeted GI screening, particularly within well-defined complexes such as the "
    "mitochondrial ribosome, where confidence score genuinely tracks GI magnitude. However, "
    "PPI data are not reliable proxies for GI data in general: the AUC of 0.630 means that "
    "a substantial fraction of high-confidence PPI pairs will not yield significant GIs, and "
    "the 223,155 significant GI pairs with no PPI counterpart represent a large, accessible "
    "space of functional relationships entirely invisible to co-purification approaches. "
    "Treating the two networks as orthogonal evidence layers in integrative models will "
    "therefore capture more functional information than treating one as a proxy for the other."
)

body(
    "Several limitations bound this analysis. The SGA dataset covers only non-essential gene "
    "pairs; the 71% of PPI pairs lacking SGA measurements are enriched for essential proteins "
    "that the current screen does not cover. Incorporating temperature-sensitive and DAmP "
    "allele data [6] would substantially increase coverage and may reveal higher overlap "
    "fractions. GI profiles are also condition-dependent [8]; interactions that are "
    "sub-threshold under standard growth conditions may reach significance under stress. "
    "Finally, co-purification efficiency is a proxy for, not a direct measure of, binding "
    "affinity; direct quantitative structural data on interface strength would enable a "
    "more mechanistic interpretation of the within-complex PPI-to-GI correlation."
)

# ── METHODS ──────────────────────────────────────────────────────────────────
head("Methods")

head("Data acquisition", level=2)
body(
    "Genetic interaction data were downloaded from thecellmap.org (Costanzo et al., 2016). "
    "The file SGA_NxN.txt contains 12,698,939 measurements for non-essential S. cerevisiae "
    "gene pairs, with columns for query and array strain identifiers, epsilon score, p-value, "
    "single-mutant fitnesses, and double-mutant fitness. Protein-protein interaction data "
    "were taken from Supplementary Data 2 of Michaelis et al. (2023), comprising a "
    "semicolon-delimited edge file (31,004 edges) and a node file (3,927 proteins), both "
    "using SGD ORF identifiers as primary keys. Markov cluster assignments from the node "
    "file (617 clusters) defined protein complex membership throughout."
)

head("Data preprocessing", level=2)
body(
    "SGA strain identifiers (format YXXXXX_sn123) were parsed to extract ORF identifiers "
    "using a regular expression matching the SGD ORF pattern. SGA data were processed in "
    "200,000-row chunks; rows in which at least one partner ORF appeared in the PPI network "
    "were retained (10,707,139 rows). Multiple SGA measurements per gene pair were aggregated "
    "by computing the mean epsilon score and minimum p-value across replicates. Gene pairs "
    "and PPI edges were represented as frozensets of sorted ORF identifiers for "
    "direction-independent matching. A GI was classified as significant if |epsilon| >= 0.08 "
    "and p < 0.05, following Costanzo et al. (2016) [2]. All analyses were performed in "
    "Python 3 using pandas, scipy.stats, numpy, scikit-learn, matplotlib, and seaborn."
)

head("Statistical analyses", level=2)
body(
    "The expected null-model overlap was computed as a binomial expectation: background_rate "
    "x n_measured_PPI_pairs, where background_rate = 223,155 / 7,721,253 = 2.90%. "
    "Enrichment significance used a normal approximation to the binomial (z-statistic). "
    "Distribution differences were assessed by two-sided Mann-Whitney U test. Within-complex "
    "versus between-complex enrichment in concordant pairs was tested by Fisher's exact test "
    "(2x2 table: [Both-within, Both-between; Nonsig-within, Nonsig-between]). GI sign "
    "enrichment relative to the genome-wide background (35.6% positive among all significant "
    "SGA pairs) was tested by Fisher's exact test with a 2x2 table of [observed-positive, "
    "observed-negative; background-positive, background-negative]. Spearman and Pearson "
    "correlations were computed with scipy.stats. Per-cluster correlations required a minimum "
    "of five concordant pairs. The ROC AUC used scikit-learn.metrics.roc_auc_score."
)

head("Data and code availability", level=2)
body(
    "All analysis scripts are available at https://github.com/hanyingjhuang/Genetic-interaction. "
    "SGA data are available at thecellmap.org (Costanzo et al., 2016). PPI data are available "
    "in Supplementary Data 2 of Michaelis et al. (2023) and at www.yeast-interactome.org."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
head("References")

for ref in [
    "1. Tong AH, Evangelista M, Parsons AB, et al. Systematic genetic analysis with ordered arrays of yeast deletion mutants. Science. 2001;294(5550):2364-2368. PMID: 11743205",
    "2. Costanzo M, VanderSluis B, Koch EN, et al. A global genetic interaction network maps a wiring diagram of cellular function. Science. 2016;353(6306):aaf1420. PMID: 27708008",
    "3. Krogan NJ, Cagney G, Yu H, et al. Global landscape of protein complexes in the yeast Saccharomyces cerevisiae. Nature. 2006;440(7084):637-643. PMID: 16554755",
    "4. Gavin AC, Aloy P, Grandi P, et al. Proteome survey reveals modularity of the yeast cell machinery. Nature. 2006;440(7084):631-636. PMID: 16429126",
    "5. Michaelis AC, Brunner AD, Zwiebel M, et al. The social and structural architecture of the yeast protein interactome. Nature. 2023;624(7990):192-200. PMID: 37968396",
    "6. Costanzo M, Baryshnikova A, Bellay J, et al. The genetic landscape of a cell. Science. 2010;327(5964):425-431. PMID: 20093466",
    "7. Collins SR, Miller KM, Maas NL, et al. Functional dissection of protein complexes involved in yeast chromosome biology using a genetic interaction map. Nature. 2007;446(7137):806-810. PMID: 17314980",
    "8. Costanzo M, Hou J, Messier V, et al. Environmental robustness of the global yeast genetic interaction network. Science. 2021;372(6542):eabf8424. PMID: 33958448",
    "9. Roguev A, Bandyopadhyay S, Zofall M, et al. Conservation and rewiring of functional modules revealed by an epistasis map in fission yeast. Science. 2008;322(5900):405-410. PMID: 18818364",
    "10. van Leeuwen J, Pons C, Mellor JC, et al. Exploring genetic suppression interactions on a global scale. Science. 2016;354(6312):aag0839. PMID: 27811238",
]:
    ref_para(ref)

# ── FIGURE LEGENDS ────────────────────────────────────────────────────────────
head("Figure Legends")

for ft, fl in [
    ("Figure 1.",
     "PPI pair coverage in the SGA dataset. Of 31,004 PPI pairs, 1,044 (3.4%, dark blue) "
     "carry a significant GI (|epsilon|>=0.08, p<0.05), 7,953 (25.6%, gray) are measured but "
     "sub-threshold, and 22,007 (71.0%, green) have no SGA measurement. Expected by chance: 261 "
     "pairs; observed: 1,044 (4.0-fold enrichment; z=49.1, p<10^-100)."),
    ("Figure 2.",
     "Absolute GI epsilon for concordant and discordant PPI pairs. Violin plots of |epsilon| for "
     "Both pairs (significant GI, dark blue; mean=0.204) and PPI-plus-non-significant pairs (gray; "
     "mean=0.027). Black lines indicate medians. Mann-Whitney U, p<10^-100."),
    ("Figure 3.",
     "PPI interaction scores across categories. Violin plots of PPI score (2-10) for Both pairs "
     "(mean=4.74), PPI-plus-nonsig pairs (mean=3.44), and PPI-only pairs. Mann-Whitney U for "
     "Both versus PPI-plus-nonsig: p=4.2x10^-49."),
    ("Figure 4.",
     "PPI score versus |GI epsilon| for concordant pairs. (Left) Scatter plot of 1,044 Both pairs "
     "colored by GI sign (blue: positive; red: negative; Spearman r=0.150, p=1.1x10^-6). "
     "(Right) Binned means with standard error bars."),
    ("Figure 5.",
     "ROC analysis for PPI score as a predictor of significant GI. Computed in the 8,997 PPI "
     "pairs with direct SGA measurement. AUC=0.630."),
    ("Figure 6.",
     "GI sign at within-complex versus between-complex interfaces. Histograms of GI epsilon for "
     "within-complex Both pairs (55.2% positive; OR=2.23 vs background, p=7.0x10^-30) and "
     "between-complex Both pairs (34.3% positive; OR=0.95, p=0.73). Dashed line at epsilon=0; "
     "red line: group mean."),
    ("Figure 7.",
     "Per-complex Spearman correlations (PPI score vs |GI epsilon|), top 15 within-complex "
     "clusters by Both-pair count (min n=5). Blue bars: p<0.05; gray bars: p>=0.05. Mitochondrial "
     "ribosome (MRPL cluster): r=0.297, p=0.010."),
    ("Figure 8.",
     "Complex-level epistasis architecture. Each bubble: one Markov cluster with >=5 Both pairs. "
     "X-axis: Spearman r (PPI score vs |GI epsilon|); y-axis: percent positive GI; bubble size: "
     "number of Both pairs; color: number of pairs (blue scale). Reference lines at r=0 and 50% "
     "positive GI. Notable complexes annotated."),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(ft + " "); font(r1, bold=True)
    r2 = p.add_run(fl); font(r2)

doc.save(OUT)
print(f"Saved: {OUT}")
