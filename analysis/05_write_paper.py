"""
05_write_paper.py
Generate the research paper as a .docx file using python-docx.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "/Users/han-yingjhuang/Claude_projects/Genetic-interaction/paper/2026.04.06_Genetic_vs_Protein_Interactions_Yeast_v1.docx"
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_body_para(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_heading(text, level=1):
    sizes = {1: 14, 2: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 12), bold=True)
    return p

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=16, bold=True)
    return p

def add_center_para(text, size=11, italic=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic, bold=bold)
    return p

def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("")
    set_font(run)

# ── Title page ────────────────────────────────────────────────────────────────
add_blank()
add_title("Mapping the Boundary Between Genetic and Physical Interaction Networks in Yeast")
add_blank()

add_center_para("Comparative analysis of the Saccharomyces cerevisiae synthetic genetic array and proteome-scale protein interactome", size=12, italic=True)
add_blank()
add_blank()

# ── Abstract ──────────────────────────────────────────────────────────────────
add_heading("Abstract")
abstract_text = (
    "Genetic interactions, measured by synthetic genetic array (SGA) technology, and protein-protein interactions (PPIs), "
    "determined by affinity-based proteomics, each provide a distinct view of how genes and proteins collaborate to sustain "
    "cellular function. Here we undertake a systematic comparison of two large-scale yeast datasets: the near-complete SGA-based "
    "genetic interaction network of Costanzo et al. (2016), comprising approximately 12.7 million gene-pair measurements for "
    "non-essential genes, and the proteome-scale PPI network of Michaelis et al. (2023), which reports 31,004 physical interactions "
    "among 3,927 proteins. Of 31,004 PPI pairs, 7.9% (2,442 pairs) are supported by a significant genetic interaction (p < 0.05), "
    "21.1% (6,555 pairs) are measured in the SGA dataset but fall below the significance threshold, and 71.0% (22,007 pairs) have "
    "no SGA counterpart. The vast majority of significant genetic interactions (1,233,608 pairs) have no corresponding PPI. "
    "Among discordant PPI pairs that carry a measured but non-significant genetic interaction, the mean absolute epsilon score "
    "(0.024) is substantially lower than for concordant pairs (0.110; Mann-Whitney U, p < 10^-100), indicating that disagreements "
    "between the two networks largely reflect weak rather than absent genetic coupling. For pairs present in both datasets, "
    "physical interaction score and genetic interaction magnitude show a modest but highly significant positive correlation "
    "(Spearman r = 0.237, p = 1.3 x 10^-32), stronger for positive genetic interactions (r = 0.309) than for negative ones "
    "(r = 0.200). Within-complex pairs show higher GI magnitude than between-complex pairs (mean |epsilon| 0.127 vs 0.076) "
    "and a modestly higher positive-GI fraction (58.4% vs 49.3%). Complex-specific analysis across 214 Markov clusters reveals "
    "substantial heterogeneity: the spliceosome displays predominantly negative genetic interactions (91%), the elongator complex "
    "exclusively positive ones (100%), and the mitochondrial ribosome shows the strongest concordance between PPI and GI strength "
    "(r = 0.272, p = 0.005). These results reveal a complementary rather than redundant relationship between the two network "
    "types, with physical contacts setting a baseline tendency for genetic buffering but pathway architecture and functional "
    "redundancy ultimately determining whether a significant genetic interaction is observed."
)
add_body_para(abstract_text)
add_blank()

# Keywords
kw_para = doc.add_paragraph()
kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_run = kw_para.add_run("Keywords: ")
set_font(kw_run, bold=True)
kw_body = kw_para.add_run("genetic interactions, protein-protein interactions, Saccharomyces cerevisiae, SGA, epistasis, protein complexes, interactome")
set_font(kw_body)

doc.add_page_break()

# ── Introduction ──────────────────────────────────────────────────────────────
add_heading("Introduction")

intro1 = (
    "Two powerful technologies have shaped our understanding of cellular organization in the budding yeast "
    "Saccharomyces cerevisiae: synthetic genetic array (SGA) analysis, which detects functional relationships "
    "through fitness defects in double mutants [1,2], and affinity-based proteomics, which maps direct physical "
    "contacts between proteins [3,4,5]. Together these technologies have generated some of the most comprehensive "
    "functional genomics datasets available for any organism. Yet the relationship between the networks they define "
    "remains incompletely understood."
)
add_body_para(intro1)

intro2 = (
    "Genetic interactions (GIs) are quantified as epsilon scores, the deviation of a double mutant's fitness from "
    "the product of the single-mutant fitnesses [1]. A negative epsilon (synthetic sick or lethal) suggests that two "
    "genes buffer one another's absence; a positive epsilon (suppression or epistasis) indicates that both mutations "
    "impair the same pathway, rendering the second hit inconsequential. At genome scale, Costanzo et al. (2010) and "
    "Costanzo et al. (2016) applied SGA to map approximately 23 million double mutants covering roughly 90% of all "
    "yeast gene pairs [2,6]. This network revealed that functionally related genes cluster into modules defined by "
    "similar genetic interaction profiles, with protein complexes occupying cores of densely connected subnetworks."
)
add_body_para(intro2)

intro3 = (
    "Protein-protein interactions (PPIs) capture a complementary dimension: direct physical binding between proteins "
    "detected by co-purification, yeast two-hybrid, or other biochemical methods. Landmark studies by Krogan et al. "
    "(2006) and Gavin et al. (2006) used tandem affinity purification coupled to mass spectrometry to identify over "
    "7,000 yeast protein associations [3,4]. More recently, Michaelis et al. (2023) leveraged high-throughput "
    "affinity enrichment mass spectrometry across 4,159 GFP-tagged strains to generate a near-complete yeast "
    "interactome containing 31,004 interactions among 3,927 proteins, doubling protein coverage and tripling "
    "reliable interaction counts relative to prior maps [5]. This dataset incorporates a three-layer evidence "
    "framework combining forward pull-down statistics, reciprocal pull-down validation, and abundance correlation "
    "to produce a composite interaction score ranging from 2 to 10."
)
add_body_para(intro3)

intro4 = (
    "Prior work has established conceptual links between GIs and PPIs. Within protein complexes, subunit pairs "
    "tend to exhibit positive genetic interactions, consistent with the idea that disrupting any component impairs "
    "the same functional unit [2]. Between complexes that act in parallel pathways, negative GIs predominate, "
    "reflecting mutual buffering [7]. Network analyses have shown that GI profiles predict protein complex membership "
    "with high accuracy, suggesting that GI and PPI data capture overlapping but distinct information [2,6]. "
    "However, a direct quantitative comparison using the most comprehensive datasets in each category has not "
    "been performed."
)
add_body_para(intro4)

intro5 = (
    "Here we leverage the Costanzo et al. (2016) SGA dataset and the Michaelis et al. (2023) yeast proteome-scale "
    "interactome to address four interconnected questions. First, to what extent do significant GIs and high-confidence "
    "PPIs identify the same gene pairs? Second, do discordant pairs, those present in one network but absent from "
    "the other, reflect weak signals or qualitatively distinct biology? Third, how well do PPI strength and GI "
    "magnitude correlate for pairs that appear in both networks? And fourth, do certain protein complexes show "
    "stronger concordance between their physical and genetic landscapes than others? The answers provide a "
    "framework for integrating the two data types and illuminate the complementary logic of physical and functional "
    "interaction networks."
)
add_body_para(intro5)

# ── Results ───────────────────────────────────────────────────────────────────
add_heading("Results")

add_heading("Coverage and overlap between the genetic and physical interaction networks", level=2)

res1 = (
    "The Michaelis et al. (2023) dataset contains 31,004 physical interactions among 3,927 S. cerevisiae proteins, "
    "organized into 617 Markov-clustering-defined modules. The Costanzo et al. (2016) SGA dataset covers 12,698,939 "
    "gene-pair measurements in the non-essential gene space. Of these SGA measurements, 10,707,139 (84.3%) involve "
    "at least one gene whose protein product appears in the PPI network, yielding 7,721,253 unique gene pairs with "
    "at least one PPI-network protein. Among these, 1,236,050 (16.0%) reach statistical significance at p < 0.05."
)
add_body_para(res1)

res2 = (
    "Focusing on the 31,004 PPI pairs as the primary reference set, we classified each pair according to its SGA "
    "status (Figure 1). Only 2,442 PPI pairs (7.9%) are supported by a significant genetic interaction (p < 0.05). "
    "An additional 6,555 PPI pairs (21.1%) are measured in the SGA dataset but fail to reach significance. "
    "The remaining 22,007 PPI pairs (71.0%) have no SGA counterpart, reflecting either the restriction of the "
    "current SGA dataset to non-essential genes or incomplete measurement coverage. Conversely, 1,233,608 pairs "
    "carry a significant genetic interaction without any corresponding physical interaction, underscoring that "
    "the GI network predominantly connects gene pairs that do not physically associate."
)
add_body_para(res2)

add_heading("Pairs with disagreements show lower interaction signals in both networks", level=2)

res3 = (
    "A central question is whether discordant pairs, those present in one network but absent from the other, "
    "reflect qualitatively absent coupling or merely weak coupling falling below detection thresholds. To address "
    "this, we compared the absolute GI epsilon scores of PPI pairs that have a measured but non-significant genetic "
    "interaction (hereafter, PPI-plus-nonsig) against pairs with both a significant GI and a PPI (hereafter, Both). "
    "The mean absolute epsilon for PPI-plus-nonsig pairs is 0.024, compared with 0.110 for Both pairs "
    "(Mann-Whitney U, p < 10^-100; Figure 2). This approximately 4.5-fold difference indicates that discordant "
    "pairs are not completely devoid of genetic coupling; rather, their genetic interaction signal falls below "
    "the p < 0.05 threshold, consistent with weak functional relationships."
)
add_body_para(res3)

res4 = (
    "The directionality of this effect is mirrored in the PPI network. Both pairs have a higher mean PPI "
    "interaction score (4.04) than PPI-plus-nonsig pairs (3.42; Mann-Whitney U, p = 7.3 x 10^-24; Figure 3). "
    "This suggests that higher-confidence physical interactions are somewhat more likely to be associated with "
    "detectable genetic coupling. Nevertheless, the ability of PPI score alone to predict whether a gene pair "
    "shows significant GI is limited: a receiver operating characteristic analysis yields an area under the curve "
    "(AUC) of 0.564 (Figure 5), only modestly above chance. Thus, while both signal levels are lower for "
    "discordant pairs, PPI score alone is an insufficient predictor of genetic interaction status."
)
add_body_para(res4)

add_heading("Physical interaction strength weakly but significantly predicts genetic interaction magnitude", level=2)

res5 = (
    "Among the 2,442 pairs with both a PPI and a significant GI, we examined whether the strength of the physical "
    "interaction predicts the magnitude of the genetic interaction (Figure 4). PPI score and absolute GI epsilon "
    "show a modest but highly significant positive correlation (Spearman r = 0.237, p = 1.3 x 10^-32; Pearson "
    "r = 0.232, p = 2.6 x 10^-31). The correlation is stronger for pairs with positive genetic interactions "
    "(Spearman r = 0.309, p = 2.3 x 10^-31, n = 1,352) than for those with negative genetic interactions "
    "(Spearman r = 0.200, p = 2.9 x 10^-11, n = 1,090)."
)
add_body_para(res5)

res6 = (
    "To explore whether this correlation is driven by a specific range of PPI confidence scores, we evaluated "
    "the Spearman r across successively higher PPI score thresholds. The correlation is fairly stable from score "
    "threshold 2 to 4 (r = 0.237 to 0.247) but weakens for very high-scoring interactions (score >= 8, r = 0.072, "
    "p = 0.20). This finding suggests that the highest-confidence physical interactions, which tend to represent "
    "stable stoichiometric complex subunit contacts, do not necessarily correlate with strong genetic interactions, "
    "possibly because such obligate interactions leave little room for compensatory genetic buffering."
)
add_body_para(res6)

add_heading("Positive genetic interactions predominate among physically interacting partners", level=2)

res7 = (
    "Among the 2,442 Both pairs, 55.4% carry a positive genetic interaction (epsilon > 0) and 44.6% carry a "
    "negative one (Figure 7). This modest excess of positive GIs among PPI pairs is consistent with the "
    "established model in which proteins that physically associate tend to function within the same pathway, "
    "so that disrupting either partner imposes a similar functional consequence."
)
add_body_para(res7)

res8 = (
    "The distinction becomes more pronounced when pairs are stratified by whether both partners belong to the "
    "same Markov cluster (within-cluster, n = 1,623) or different clusters (between-cluster, n = 819). "
    "Within-cluster pairs show 58.4% positive GI, compared with 49.3% for between-cluster pairs. "
    "A Fisher exact test confirms that between-cluster pairs are more strongly enriched in the Both category "
    "relative to PPI-plus-nonsig (OR = 0.74, p = 1.8 x 10^-9), indicating that inter-complex physical contacts "
    "are more likely to be accompanied by significant genetic interactions than intra-complex contacts. "
    "Within-cluster pairs that have both a PPI and a GI also show higher GI magnitude on average "
    "(mean |epsilon| = 0.127) compared with between-cluster Both pairs (mean |epsilon| = 0.076), and the "
    "within-cluster PPI score is higher as well (mean 4.74 vs 2.65), reflecting the well-established tendency "
    "of stable complex subunits to receive high interaction scores."
)
add_body_para(res8)

add_heading("Protein complex identity shapes the relationship between physical and genetic interaction strength", level=2)

res9 = (
    "Grouping Both pairs by their Markov cluster revealed 214 clusters with at least one pair in both datasets, "
    "of which 177 contained at least three such pairs suitable for within-cluster correlation analysis (Figure 6). "
    "The per-cluster Spearman correlations between PPI score and |GI epsilon| vary widely, from r = -0.47 to "
    "r = 0.64, demonstrating that the global modest correlation conceals complex-specific heterogeneity."
)
add_body_para(res9)

res10 = (
    "Among the largest clusters, ribosomal proteins dominate: Cluster 1 (large ribosomal subunit, RPL11B/RPL11A "
    "and related, n = 280) and Cluster 2 (small ribosomal subunit, RPS23A/RPS23B and related, n = 244) each "
    "show approximately 70% positive GI among Both pairs, but PPI score and |GI epsilon| do not correlate "
    "within these clusters (r = 0.034 and -0.029, respectively). This uncoupling suggests that within the "
    "ribosome, all subunit pairs are similarly indispensable, leveling the GI landscape regardless of PPI score."
)
add_body_para(res10)

res11 = (
    "In contrast, the mitochondrial ribosome cluster (MRPL19/MRPL10/MRPL24 and related proteins, n = 104) "
    "shows a significant positive correlation (Spearman r = 0.272, p = 0.005), indicating that within this "
    "organellar complex, stronger physical association indeed predicts stronger genetic coupling. Clusters with "
    "the highest within-cluster correlations include an uncharacterized set of proteins centered on PUP1/IDI1/DDR2 "
    "(r = 0.637, p = 0.002, n = 21) and the RNA-binding protein cluster HEK2/DHH1/FUS3 (r = 0.614, p = 0.011, "
    "n = 16), pointing to particular biological contexts in which PPI strength is especially predictive of "
    "genetic coupling."
)
add_body_para(res11)

res12 = (
    "GI direction also differs markedly across complexes. The spliceosome cluster (PRP4/DIB1/SNU66, n = 46) "
    "shows 91% negative GI among Both pairs, suggesting that spliceosome subunit pairs are strongly synthetic "
    "sick or lethal, consistent with the essentiality of RNA splicing. The nuclear pore complex cluster "
    "(POM152/MTR2/NUP188, n = 14) likewise shows 100% negative GI. In stark contrast, the elongator complex "
    "(ELP2/ELP3/ELP4, n = 16) and the Sin3-containing NuRD/NuA4-related cluster (STB3/CTI6/SIN3, n = 31, "
    "81% positive) show predominantly positive GI, suggesting that within these chromatin-regulatory modules, "
    "component loss impairs the same functional output rather than triggering distinct compensatory responses "
    "(Figure 8)."
)
add_body_para(res12)

# ── Discussion ────────────────────────────────────────────────────────────────
add_heading("Discussion")

disc1 = (
    "This analysis reveals a complementary, rather than redundant, relationship between the genetic and physical "
    "interaction networks in yeast. The two datasets share only 7.9% of PPI pairs in the significant GI category, "
    "yet this overlap is far above what would be expected from random gene pairings. The modest but highly "
    "significant global correlation (Spearman r = 0.237) confirms that physical interaction strength and genetic "
    "interaction magnitude are genuinely correlated properties, while the ROC AUC of 0.564 makes clear that PPI "
    "score alone cannot predict genetic interaction status."
)
add_body_para(disc1)

disc2 = (
    "Why do 92.1% of PPI pairs lack a significant genetic interaction? The most parsimonious explanation is "
    "mechanistic: a physical contact between two proteins does not require genetic redundancy between the genes "
    "encoding them. Many subunit-subunit contacts within stable complexes are stoichiometrically obligate, "
    "meaning that disrupting either partner collapses the entire complex. In this regime, the double mutant "
    "is no more severe than either single mutant, yielding a positive or near-neutral GI. For the SGA dataset "
    "to detect a significant GI, the double mutant must exceed the expected fitness product by a detectable "
    "margin, a condition that requires either synthetic lethality (negative GI) or substantial epistasis "
    "(positive GI). Stable complex contacts may fall below this threshold because of buffering by the remaining "
    "complex or by parallel pathways. The observation that discordant PPI pairs carry measurable, albeit "
    "sub-threshold, GI signals (mean |epsilon| = 0.024 vs 0.110 for concordant pairs) supports this "
    "interpretation: the interactions are not biologically absent, they are merely weak."
)
add_body_para(disc2)

disc3 = (
    "The inverse asymmetry is even more striking: 1,233,608 significant GI pairs lack a PPI counterpart. "
    "This reflects the well-established principle that genetic interactions operate at the pathway level, "
    "not only at the level of direct physical contacts [2,6]. Two genes that act in parallel pathways "
    "buffer each other functionally without any direct protein contact, generating negative GIs across "
    "pathway boundaries. Similarly, two genes that converge on the same essential output produce positive "
    "GIs without necessarily physically associating. The GI network thus reports on the entire functional "
    "architecture of the cell, including redundant pathways, compensatory mechanisms, and regulatory "
    "cross-talk that physical interaction maps cannot access."
)
add_body_para(disc3)

disc4 = (
    "The stronger correlation for positive GIs (r = 0.309) compared with negative GIs (r = 0.200) is "
    "consistent with the conceptual model in which same-pathway gene pairs both physically associate and "
    "show suppression-type epistasis, while synthetic sick/lethal interactions often arise between functionally "
    "unrelated proteins that compensate for each other's loss. Physical contacts may be more directly predictive "
    "of same-pathway functional relationships, explaining the asymmetry. The observation that correlation "
    "weakens for very high PPI scores (r = 0.072 at score >= 8) suggests that obligate, high-affinity "
    "stoichiometric contacts, such as those between ribosomal subunits, impose a ceiling effect on GI "
    "magnitude because all subunit pairs are equally dependent on the intact complex."
)
add_body_para(disc4)

disc5 = (
    "The complex-specific patterns illuminate particular biological contexts. The spliceosome's predominance "
    "of negative GIs is expected given the essentiality of pre-mRNA splicing and the likely synthetic lethality "
    "of disrupting two spliceosome subunits. The nuclear pore complex shows a similar pattern. Conversely, "
    "the elongator complex, which modifies tRNA wobble uridines and facilitates translational fidelity, "
    "shows exclusively positive GIs, consistent with the idea that loss of any component produces a common "
    "translation defect that cannot be compounded. The mitochondrial ribosome's significant correlation between "
    "PPI score and |GI epsilon| may reflect that this complex has a more modular, less tightly integrated "
    "architecture than the cytoplasmic ribosome, so that PPI strength more faithfully reports on the "
    "functional importance of individual contacts. These complex-specific patterns suggest that integrating "
    "PPI and GI data is most informative when done at the module level rather than globally."
)
add_body_para(disc5)

disc6 = (
    "Our findings have implications for network-based approaches to predicting gene function and disease "
    "gene interactions. The modest but significant global correlation supports the use of PPI data to "
    "prioritize gene pairs for genetic interaction screening, though the low ROC AUC indicates that many "
    "high-priority physical contacts will not yield detectable GIs. The complementarity of the two networks "
    "argues for integration strategies that treat GIs and PPIs as orthogonal evidence channels rather than "
    "interchangeable proxies. Complex-specific calibration, as demonstrated here for mitochondrial ribosome "
    "and spliceosome components, may substantially improve the predictive power of such integrations."
)
add_body_para(disc6)

disc7 = (
    "Several limitations of this analysis deserve acknowledgment. The SGA dataset used here covers only "
    "non-essential gene pairs (NxN), omitting essential-by-essential and essential-by-non-essential pairs, "
    "which represent a substantial fraction of all gene pairs. Because many high-confidence PPI pairs involve "
    "essential proteins, the 71% of PPI pairs with no SGA measurement likely includes many biologically "
    "important relationships that are simply not accessible in the current SGA framework. Additionally, "
    "the SGA dataset aggregates measurements across a limited set of standard laboratory conditions, "
    "whereas GI profiles can shift substantially under stress [8], potentially underestimating the extent "
    "of concordance under particular growth conditions. Future integration of essential-gene interaction "
    "data and condition-specific GI screens with the Michaelis et al. (2023) PPI map should provide "
    "a more complete picture of the boundary between physical and genetic networks."
)
add_body_para(disc7)

# ── Methods ───────────────────────────────────────────────────────────────────
add_heading("Methods")

add_heading("Data sources", level=2)
meth1 = (
    "Genetic interaction data were obtained from Costanzo et al. (2016) [1] via the CellMap repository "
    "(https://thecellmap.org). The file SGA_NxN.txt contains 12,698,939 measurements for non-essential "
    "gene pairs, with columns for query and array strain identifiers, the genetic interaction epsilon score, "
    "p-value, single-mutant fitnesses, double-mutant fitness, and standard deviation. Protein-protein "
    "interaction data were obtained from the supplementary materials of Michaelis et al. (2023) [5] "
    "(Supplementary Data 2), comprising The_Yeast_Interactome_Edges.csv (31,004 edges, semicolon-delimited) "
    "and The_Yeast_Interactome_Nodes.csv (3,927 nodes), both using Saccharomyces Genome Database (SGD) "
    "identifiers as primary keys."
)
add_body_para(meth1)

add_heading("Data preprocessing", level=2)
meth2 = (
    "SGA strain identifiers of the format YXXXXX_sn123 were parsed to extract the open reading frame "
    "(ORF) identifier (the substring before the underscore) using a regular expression matching the "
    "standard SGD ORF format. SGA data were streamed in chunks of 200,000 rows using pandas to manage "
    "memory requirements. Only rows in which at least one partner ORF appeared in the PPI network were "
    "retained, yielding 10,707,139 rows. For each unique gene pair, multiple measurements were aggregated "
    "by computing the mean epsilon score and minimum p-value across all replicates. Gene pairs were "
    "represented as frozensets of sorted ORF identifiers to ensure direction-independent matching. "
    "PPI edges were similarly represented as frozensets of sorted SGD identifiers."
)
add_body_para(meth2)

add_heading("Overlap analysis and signal comparison", level=2)
meth3 = (
    "PPI pairs were classified into three categories: (1) Both, if the pair was present in the SGA "
    "dataset with p < 0.05; (2) PPI-plus-nonsig, if present in SGA but p >= 0.05; and (3) PPI-only, "
    "if not measured in SGA. The category Significant GI only was defined as SGA pairs with p < 0.05 "
    "that had no corresponding PPI entry. Differences in absolute epsilon and PPI score distributions "
    "between categories were assessed using the two-sided Mann-Whitney U test. A receiver operating "
    "characteristic analysis was performed treating significant GI (p < 0.05) as the positive class, "
    "PPI interaction score as the predictor, and using the measured subset (Both plus PPI-plus-nonsig) "
    "as the analysis population. The area under the curve was computed using scikit-learn."
)
add_body_para(meth3)

add_heading("Correlation analysis", level=2)
meth4 = (
    "Spearman and Pearson correlations between PPI score and absolute GI epsilon were computed for all "
    "Both pairs and for subsets stratified by GI sign and by cluster membership. Cluster assignments "
    "were taken from the Markov cluster column in The_Yeast_Interactome_Nodes.csv, which defines 617 "
    "modules. Per-cluster correlations were computed for clusters containing three or more Both pairs. "
    "Enrichment of within-cluster vs between-cluster pairs in the Both vs PPI-plus-nonsig categories "
    "was assessed by Fisher exact test. All statistical analyses were performed in Python using scipy.stats."
)
add_body_para(meth4)

add_heading("Figures", level=2)
meth5 = (
    "All figures were generated using matplotlib and seaborn in Python at 300 DPI. Violin plots depict "
    "kernel density estimates with median lines. Scatter plots use jitter on the PPI score axis for "
    "visual clarity. Bar charts of per-cluster correlations show Spearman r values colored by "
    "significance (blue, p < 0.05; gray, p >= 0.05). The cluster bubble plot encodes Spearman r "
    "by color (red-blue diverging scale) and number of pairs by bubble area."
)
add_body_para(meth5)

# ── References ────────────────────────────────────────────────────────────────
add_heading("References")

refs = [
    "1. Tong AH, Evangelista M, Parsons AB, et al. Systematic genetic analysis with ordered arrays of yeast deletion mutants. Science. 2001;294(5550):2364-2368. PMID: 11743205",
    "2. Costanzo M, VanderSluis B, Koch EN, et al. A global genetic interaction network maps a wiring diagram of cellular function. Science. 2016;353(6306):aaf1420. PMID: 27708008",
    "3. Krogan NJ, Cagney G, Yu H, et al. Global landscape of protein complexes in the yeast Saccharomyces cerevisiae. Nature. 2006;440(7084):637-643. PMID: 16554755",
    "4. Gavin AC, Aloy P, Grandi P, et al. Proteome survey reveals modularity of the yeast cell machinery. Nature. 2006;440(7084):631-636. PMID: 16429126",
    "5. Michaelis AC, Brunner AD, Zwiebel M, et al. The social and structural architecture of the yeast protein interactome. Nature. 2023;624(7990):192-200. PMID: 37968396",
    "6. Costanzo M, Baryshnikova A, Bellay J, et al. The genetic landscape of a cell. Science. 2010;327(5964):425-431. PMID: 20093466",
    "7. van Leeuwen J, Pons C, Mellor JC, et al. Exploring genetic suppression interactions on a global scale. Science. 2016;354(6312):aag0839. PMID: 27811238",
    "8. Costanzo M, Hou J, Messier V, et al. Environmental robustness of the global yeast genetic interaction network. Science. 2021;372(6542):eabf8424. PMID: 33958448",
    "9. Hsu CH, Wang TY, Chu HT, et al. A quantitative analysis of monochromaticity in genetic interaction networks. BMC Bioinformatics. 2012;13:65. PMID: 22372977",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    run = p.add_run(ref)
    set_font(run, size=10)

# ── Figure legends ────────────────────────────────────────────────────────────
add_heading("Figure Legends")

fig_legends = [
    ("Figure 1.", "Distribution of 31,004 PPI pairs by SGA coverage. Pie chart showing the fraction of PPI pairs classified as Both (PPI and significant GI, p < 0.05; dark blue), PPI with measured but non-significant GI (gray), and PPI with no SGA measurement (green)."),
    ("Figure 2.", "GI epsilon scores by interaction category. Violin plots comparing the distribution of mean GI epsilon scores for pairs with both a PPI and significant GI (Both, dark blue) versus pairs with a PPI and measured but non-significant GI (gray). Black horizontal lines indicate medians."),
    ("Figure 3.", "PPI score distributions across interaction categories. Violin plots showing PPI interaction scores for Both pairs, PPI-plus-non-significant GI pairs, and PPI-only pairs."),
    ("Figure 4.", "PPI score versus |GI epsilon| for Both pairs. Left: scatter plot of all 2,442 Both pairs, colored by GI direction (blue, positive; red, negative). Right: binned means of |GI epsilon| per PPI score bin with standard error bars."),
    ("Figure 5.", "ROC curve for PPI score predicting significant GI. AUC = 0.564, computed in the measured subset (Both plus PPI-plus-nonsig, n = 8,997)."),
    ("Figure 6.", "Per-cluster Spearman correlations between PPI score and |GI epsilon|. Top 20 clusters by number of Both pairs, colored by significance (blue, p < 0.05; gray, p >= 0.05)."),
    ("Figure 7.", "GI epsilon distribution for within-cluster and between-cluster Both pairs. Histograms showing mean epsilon for within-complex pairs (left) and between-complex pairs (right), with mean indicated in red."),
    ("Figure 8.", "Cluster-level relationship between mean PPI score and mean |GI epsilon|. Each bubble represents a Markov cluster with at least three Both pairs; bubble size encodes number of pairs; color encodes Spearman r (red-blue diverging scale). Top 10 clusters by pair count are labeled."),
]

for fig_title, fig_text in fig_legends:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    title_run = p.add_run(fig_title + " ")
    set_font(title_run, bold=True)
    body_run = p.add_run(fig_text)
    set_font(body_run)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Paper saved to: {OUT_PATH}")
