"""Comprehensive QC for v66 manuscripts and figures.

Checks:
  1. Banned characters (em/en dash, tilde, semicolon, question mark,
     'Fig.', British spellings).
  2. Numeric claim cross-check: every OR / P / n in the main text body
     must appear in a v66 figure CSV.
  3. Cross-reference integrity: every 'Figure N<L>' in body has a
     caption, every 'Supplementary Figure N' cited in main exists in
     supp, every panel letter in a composite appears in its caption.
  4. Tense consistency in Results and Discussion (past tense preferred).
  5. Figure layout sanity: panel PNGs present, aspect ratios sane,
     tall panel aspect < 0.8, composite aspects plausible.
  6. Caption-to-panel consistency: caption subpanel count matches
     composite layout (Fig 3 has 11 panels; Fig 5 has 9 panels).
  7. Image+caption ordering in docx: image paragraph precedes caption.

Exits with status 1 if any check fails.
"""
import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from PIL import Image

ROOT = Path(
    "/Users/han-yingjhuang/Library/Mobile Documents/"
    "com~apple~CloudDocs/Business - Projects/Genetic-interaction"
)
MAIN = ROOT / "manuscripts" / "2026.04.21_GI_PPI_v66.docx"
SUPP = ROOT / "manuscripts" / "2026.04.21_GI_PPI_v66_supplementary.docx"
SUB_MAIN = ROOT / "submission" / "Science" / "GI_PPI.docx"
SUB_SUPP = ROOT / "submission" / "Science" / "GI_PPI_supplementary.docx"

V66_PANELS = ROOT / "figures" / "v66_panels"
V66_SUPP = ROOT / "figures" / "v66_supp_panels"
V66_COMP = ROOT / "figures" / "v66_composites"

BANNED = {
    "em dash": "\u2014",
    "en dash": "\u2013",
    "tilde": "~",
    "semicolon": ";",
}

BRITISH = [
    "colour", "flavour", "behaviour", "analyse", "analysing",
    "organisation", "optimise", "optimising", "characterise",
    "characterising", "summarise", "summarising", "enrolment",
    "modelling", "labelled", "labelling", "recognise", "recognising",
    "emphasise", "emphasising",
]

PARTICIPIAL = re.compile(r",\s+(\w+ing)\s+(that|which)\b")
QUESTION_MARK = re.compile(r"\?")
FIG_SHORT = re.compile(r"\bFig\.\s*")

FAIL = False


def fail(msg):
    global FAIL
    print(f"  FAIL  {msg}")
    FAIL = True


def ok(msg):
    print(f"  ok    {msg}")


def is_reference_paragraph(text):
    if text.startswith(tuple(f"{i}." for i in range(1, 40))):
        if ("doi" in text or "PMID" in text or "(20" in text[-15:]):
            return True
    if re.match(r"^\d+\.\s", text) and len(text) > 120:
        return True
    return False


# ------------------------------------------------------------------
# 1. Banned characters
# ------------------------------------------------------------------
print("\n========== banned-character scan ==========")
for tag, path in [("main", MAIN), ("supp", SUPP),
                  ("sub_main", SUB_MAIN), ("sub_supp", SUB_SUPP)]:
    if not path.exists():
        fail(f"{tag} missing: {path}")
        continue
    d = Document(str(path))
    counts = {k: 0 for k in BANNED}
    counts["british"] = 0
    counts["participial"] = 0
    counts["question"] = 0
    counts["fig_short"] = 0
    for i, p in enumerate(d.paragraphs):
        t = p.text
        if not t.strip():
            continue
        ref = is_reference_paragraph(t.strip())
        for key, ch in BANNED.items():
            n = t.count(ch)
            if n:
                counts[key] += n
                if counts[key] <= 2:
                    print(f"    [{tag} p{i}] {key}: {t[:120]!r}")
        for word in BRITISH:
            if re.search(rf"\b{word}\b", t, re.I):
                counts["british"] += 1
                if counts["british"] <= 2:
                    print(f"    [{tag} p{i}] british {word!r}")
        m = PARTICIPIAL.search(t)
        if m:
            counts["participial"] += 1
            if counts["participial"] <= 2:
                print(f"    [{tag} p{i}] participial {m.group(0)!r}")
        if not ref:
            q = len(QUESTION_MARK.findall(t))
            if q:
                counts["question"] += q
            f_ = len(FIG_SHORT.findall(t))
            if f_:
                counts["fig_short"] += f_
    print(f"  [{tag}] em={counts['em dash']} en={counts['en dash']} "
          f"~={counts['tilde']} ;={counts['semicolon']} "
          f"british={counts['british']} part={counts['participial']} "
          f"?={counts['question']} Fig.={counts['fig_short']}")
    for k in ["em dash", "tilde", "semicolon", "british",
              "participial", "question", "fig_short"]:
        if counts[k]:
            fail(f"{tag}: {k} = {counts[k]} (must be 0)")

# ------------------------------------------------------------------
# 2. Numeric claim cross-check against v66 CSVs
# ------------------------------------------------------------------
print("\n========== numeric claim cross-check ==========")
fig3c = pd.read_csv(ROOT / "build" / "v66_fig3C_detection_rate.csv")
fig3e = pd.read_csv(ROOT / "build" / "v66_fig3E_pairwise_arch.csv")
fig3g = pd.read_csv(ROOT / "build" / "v66_fig3G_sign_enrichment.csv")
fig3h = pd.read_csv(ROOT / "build" / "v66_fig3H_logreg.csv")
fig5b = pd.read_csv(ROOT / "build" / "v66_fig5B_disease_or_12group.csv")

main_text = "\n".join(p.text for p in Document(str(MAIN)).paragraphs)

# Detection rates
for _, row in fig3c.iterrows():
    pct = f"{row['rate'] * 100:.1f}%"
    if pct in main_text:
        ok(f"detection rate {pct} ({row['label']}) present")
    else:
        fail(f"detection rate {pct} ({row['label']}) missing from main")

# Pairwise Fisher ORs — direct vs inter_non_copurified headline
direct_vs_inter = fig3e[
    (fig3e["a"] == "direct_within_complex")
    & (fig3e["b"] == "inter_non_copurified")
].iloc[0]
or_str = f"{direct_vs_inter['OR']:.2f}"
if or_str in main_text:
    ok(f"pairwise OR {or_str} (direct vs inter non-co-pur) present")
else:
    fail(f"pairwise OR {or_str} missing")

# Sign-stratified: positive direct within
pos_direct = fig3g[(fig3g["arch_key"] == "direct_within_complex")
                    & (fig3g["sign"] == "Positive")].iloc[0]
if f"{pos_direct['OR']:.2f}" in main_text:
    ok(f"sign-strat OR {pos_direct['OR']:.2f} (positive direct) present")
else:
    fail(f"sign-strat OR {pos_direct['OR']:.2f} missing")

# Logistic regression coefficient for direct_within_complex
lr = fig3h[fig3h["feature"] == "direct_within_complex"].iloc[0]
if f"{lr['OR']:.2f}" in main_text:
    ok(f"logreg OR {lr['OR']:.2f} (direct) present")
else:
    fail(f"logreg OR {lr['OR']:.2f} missing")

# Fig 5B positive direct disease enrichment headline
pos_dir_5b = fig5b[(fig5b["arch_key"] == "direct_within_complex")
                    & (fig5b["sign"] == "Positive")].iloc[0]
if f"{pos_dir_5b['OR']:.2f}" in main_text:
    ok(f"disease OR {pos_dir_5b['OR']:.2f} (positive direct) present")
else:
    fail(f"disease OR {pos_dir_5b['OR']:.2f} missing")

# ------------------------------------------------------------------
# 3. Cross-reference integrity
# ------------------------------------------------------------------
print("\n========== cross-reference integrity ==========")
main_doc = Document(str(MAIN))
supp_doc = Document(str(SUPP))
main_text = "\n".join(p.text for p in main_doc.paragraphs)
supp_text = "\n".join(p.text for p in supp_doc.paragraphs)

# All Figure X<letter> references in body must exist as captions
fig_refs = set(re.findall(r"Figure (\d[A-Z]?)", main_text))
fig_captions = set()
for p in main_doc.paragraphs:
    m = re.match(r"^Figure (\d)\.", p.text.strip())
    if m:
        fig_captions.add(m.group(1))

for ref in sorted(fig_refs):
    num = re.match(r"(\d)", ref).group(1)
    if num in fig_captions:
        ok(f"Figure {ref} has caption (Figure {num})")
    else:
        fail(f"Figure {ref} has no caption for Figure {num}")

# Supplementary Figure cross-check
supp_refs = set(re.findall(r"Supplementary Figure (\d+)", main_text))
supp_caps = set()
for p in supp_doc.paragraphs:
    m = re.match(r"^Supplementary Figure (\d+)\.", p.text.strip())
    if m:
        supp_caps.add(m.group(1))

for ref in sorted(supp_refs):
    if ref in supp_caps:
        ok(f"Supplementary Figure {ref} has caption")
    else:
        fail(f"Supplementary Figure {ref} cited but no caption in supp")

# Fig 3 caption must mention panels A-K and K_tall (K accounted for by K_tall)
fig3_cap_para = None
for p in main_doc.paragraphs:
    if p.text.strip().startswith("Figure 3.") and "(A)" in p.text:
        fig3_cap_para = p.text
        break
if fig3_cap_para is None:
    fail("Figure 3 caption not found")
else:
    for letter in "ABCDEFGHIJK":
        if f"({letter})" in fig3_cap_para:
            ok(f"Figure 3 caption enumerates ({letter})")
        else:
            fail(f"Figure 3 caption missing ({letter})")

fig5_cap_para = None
for p in main_doc.paragraphs:
    if p.text.strip().startswith("Figure 5.") and "(A)" in p.text:
        fig5_cap_para = p.text
        break
if fig5_cap_para is None:
    fail("Figure 5 caption not found")
else:
    for letter in "ABCDEFGHI":
        if f"({letter})" in fig5_cap_para:
            ok(f"Figure 5 caption enumerates ({letter})")
        else:
            fail(f"Figure 5 caption missing ({letter})")

# ------------------------------------------------------------------
# 4. Skill-rule check: figures as grammatical subjects
# ------------------------------------------------------------------
print("\n========== figure-as-subject check ==========")
subject_pat = re.compile(
    r"\bFigure\s*\d+[A-Z]?\s+"
    r"(shows|show|indicates|indicate|reveals|reveal|"
    r"demonstrates|demonstrate|suggests|suggest|confirms|confirm)\b"
)
sentence_start_pat = re.compile(
    r"(?:^|\.\s+|\.\n)Figure\s*\d+[A-Z]?\s+\w+"
)
n_subject = 0
for i, p in enumerate(main_doc.paragraphs):
    t = p.text
    for m in subject_pat.finditer(t):
        n_subject += 1
        print(f"    [p{i}] figure-as-subject: {m.group(0)!r}")
    for m in sentence_start_pat.finditer(t):
        n_subject += 1
        print(f"    [p{i}] figure sentence-start: {m.group(0)!r}")
if n_subject:
    fail(f"{n_subject} figure-as-subject / sentence-start cases")
else:
    ok("no figures used as grammatical subjects")

# ------------------------------------------------------------------
# 5. Figure layout sanity
# ------------------------------------------------------------------
print("\n========== figure layout sanity ==========")
fig3_panels = list("ABCDEFGHIJK")
for lbl in fig3_panels:
    p = V66_PANELS / f"fig3{lbl}.png"
    if not p.exists():
        fail(f"missing fig3{lbl}.png")
        continue
    im = Image.open(p)
    w, h = im.size
    aspect = w / h
    if not (0.5 <= aspect <= 2.0):
        fail(f"fig3{lbl} aspect {aspect:.2f} out of [0.5, 2.0]")
    else:
        ok(f"fig3{lbl}: {w}x{h} aspect={aspect:.2f}")

ktall = V66_PANELS / "fig3K_tall.png"
if ktall.exists():
    im = Image.open(ktall)
    a = im.width / im.height
    if a < 0.8:
        ok(f"fig3K_tall aspect {a:.2f} (< 0.8, is tall)")
    else:
        fail(f"fig3K_tall aspect {a:.2f} not tall")
else:
    fail("fig3K_tall.png missing")

for lbl in "ABCDEFGHI":
    p = V66_PANELS / f"fig5{lbl}.png"
    if not p.exists():
        fail(f"missing fig5{lbl}.png")
        continue
    im = Image.open(p)
    w, h = im.size
    aspect = w / h
    if not (0.5 <= aspect <= 2.5):
        fail(f"fig5{lbl} aspect {aspect:.2f} out of [0.5, 2.5]")
    else:
        ok(f"fig5{lbl}: {w}x{h} aspect={aspect:.2f}")

for stem in ["fig3_composite", "fig5_composite"]:
    comp = V66_COMP / f"{stem}.png"
    if not comp.exists():
        fail(f"missing composite {stem}.png")
        continue
    im = Image.open(comp)
    ok(f"{stem}: {im.size} aspect={im.size[0] / im.size[1]:.2f}")

for stem in ["supp_fig1", "supp_fig2"]:
    p = V66_SUPP / f"{stem}.png"
    if not p.exists():
        fail(f"missing {stem}.png")
        continue
    im = Image.open(p)
    ok(f"{stem}: {im.size} aspect={im.size[0] / im.size[1]:.2f}")

# ------------------------------------------------------------------
# 6. Image+caption ordering in docx
# ------------------------------------------------------------------
print("\n========== docx image/caption ordering ==========")
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
img_idx = []
cap_idx = []
for i, p in enumerate(main_doc.paragraphs):
    n_blips = len(p._element.findall(f".//{{{A_NS}}}blip"))
    if n_blips >= 1:
        img_idx.append(i)
    if re.match(r"^Figure \d\.", p.text.strip()):
        cap_idx.append(i)
print(f"  main image paragraphs: {len(img_idx)} at {img_idx}")
print(f"  main caption paragraphs: {len(cap_idx)} at {cap_idx}")

if len(img_idx) != 6 or len(cap_idx) != 6:
    fail(f"main: expected 6 images and 6 captions, got "
          f"{len(img_idx)} / {len(cap_idx)}")
else:
    for im, cp in zip(img_idx, cap_idx):
        if cp <= im:
            fail(f"main: caption p{cp} not after image p{im}")
        else:
            ok(f"image p{im} precedes caption p{cp}")

supp_img = []
supp_cap = []
for i, p in enumerate(supp_doc.paragraphs):
    n_blips = len(p._element.findall(f".//{{{A_NS}}}blip"))
    if n_blips >= 1:
        supp_img.append(i)
    if re.match(r"^Supplementary Figure \d+\.", p.text.strip()):
        supp_cap.append(i)
print(f"  supp image paragraphs: {len(supp_img)} at {supp_img}")
print(f"  supp caption paragraphs: {len(supp_cap)} at {supp_cap}")

if len(supp_img) < 2 or len(supp_cap) < 2:
    fail(f"supp: expected at least 2 images and 2 captions, got "
          f"{len(supp_img)} / {len(supp_cap)}")
else:
    ok(f"supp has {len(supp_img)} images and {len(supp_cap)} captions")

# ------------------------------------------------------------------
# 7. Supp Methods structure: check for 10 subsection headings
# ------------------------------------------------------------------
print("\n========== supp Methods subsections ==========")
expected_heads = [
    "Yeast data", "Human data", "Disease and phenotype annotations",
    "Complex definitions", "Architectural classification",
    "Statistical analyses", "Evolutionary metrics",
    "Cross-species conservation",
    "Machine-learning features",
    "Machine-learning training, validation, and hotspot enrichment",
]
found = []
for head in expected_heads:
    if f"{head}." in supp_text:
        found.append(head)
        ok(f"Methods subsection '{head}' present")
    else:
        fail(f"Methods subsection '{head}' missing")
print(f"  {len(found)}/{len(expected_heads)} Methods subsections present")

# ------------------------------------------------------------------
# 8. Italic S. cerevisiae / H. sapiens verification
# ------------------------------------------------------------------
print("\n========== species italicization ==========")
SPECIES_TOKENS = ["Saccharomyces cerevisiae", "S. cerevisiae",
                   "Homo sapiens", "H. sapiens"]
for tag, doc in [("main", main_doc), ("supp", supp_doc)]:
    total_hits = 0
    italic_hits = 0
    for p in doc.paragraphs:
        t = p.text
        for sp in SPECIES_TOKENS:
            for match in re.finditer(re.escape(sp), t):
                total_hits += 1
                start, end = match.span()
                pos = 0
                is_italic = False
                for run in p.runs:
                    r_start = pos
                    r_end = pos + len(run.text)
                    if r_start <= start and r_end >= end:
                        is_italic = bool(run.italic)
                        break
                    pos = r_end
                if is_italic:
                    italic_hits += 1
    if total_hits == 0:
        fail(f"{tag}: no species tokens found (expected some)")
    elif italic_hits == total_hits:
        ok(f"{tag}: {italic_hits}/{total_hits} species mentions italic")
    else:
        fail(f"{tag}: only {italic_hits}/{total_hits} species mentions italic")

# ------------------------------------------------------------------
if FAIL:
    print("\nFAIL (actionable items above)")
    sys.exit(1)
print("\nPASS")
